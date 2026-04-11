"""Generate bilingual NCS-style Word document for HPR surrogate paper."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # also set East-Asian font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.insert(0, rFonts)

def add_para(doc, text, style='Normal', bold=False, size=11,
             italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
             color=None, keep_with_next=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = keep_with_next
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading(doc, text, level=1):
    """Add a styled heading paragraph."""
    sizes = {1: 14, 2: 12, 3: 11}
    p = add_para(doc, text, bold=True, size=sizes.get(level, 11),
                 space_before=12, space_after=4)
    return p

def add_bilingual(doc, en_text, cn_text, size=11):
    """Add English paragraph then Chinese paragraph."""
    add_para(doc, en_text, size=size, space_before=0, space_after=3)
    add_para(doc, cn_text, size=size, space_before=0, space_after=6,
             color=(70, 70, 70))

def add_fig_placeholder(doc, label, caption_en, caption_cn):
    p = add_para(doc, f"[{label}]", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=8, space_after=2, color=(120, 120, 120))
    add_para(doc, f"Figure placeholder (SVG). {caption_en}", size=9,
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2,
             color=(100, 100, 100))
    add_para(doc, caption_cn, size=9, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8,
             color=(100, 100, 100))

def add_table_placeholder(doc, label, caption_en, caption_cn):
    p = add_para(doc, f"[{label}]", size=10, bold=True,
                 space_before=8, space_after=2, color=(60, 60, 120))
    add_para(doc, f"Table placeholder. {caption_en}", size=9,
             italic=True, space_before=0, space_after=2, color=(80, 80, 80))
    add_para(doc, caption_cn, size=9, italic=True,
             space_before=0, space_after=8, color=(80, 80, 80))

def hr(doc):
    """Thin horizontal rule as a paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

# ── build document ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title ─────────────────────────────────────────────────────────────────────
add_para(doc,
    "Probabilistic Neural Surrogates for Uncertainty-to-Risk Analysis\n"
    "in Coupled Multi-Physics Systems: Application to a Heat-Pipe-Cooled Reactor",
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=6, space_after=4)

add_para(doc,
    "概率神经代理模型在耦合多物理场系统不确定性–风险分析中的应用\n——以热管冷却反应堆为例",
    bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
    space_before=0, space_after=8, color=(50, 50, 50))

# ── Authors ───────────────────────────────────────────────────────────────────
add_para(doc,
    "Yinuo Chen\u00b9\u1d43\u02b8\u1d47, Sicheng Wang\u00b2\u1d47\u02b8\u1d9c, "
    "Xiaojing Liu\u00b3\u1d43\u02b8\u1d47 and Tengfei Zhang\u2074\u1d43\u02b8\u1d47\u002c*",
    size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)

# Affiliations
affiliations = [
    "\u1d43  School of Nuclear Science and Engineering, Shanghai Jiao Tong University, Shanghai, 200240, China",
    "\u1d47  Shanghai Digital Nuclear Reactor Technology Integration Innovation Center, Shanghai, 200240, China",
    "\u1d9c  College of Smart Energy, Shanghai Jiao Tong University, Shanghai, 200240, China",
    "*  Corresponding author: zhangtengfei@sjtu.edu.cn",
]
for aff in affiliations:
    add_para(doc, aff, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=2, color=(80, 80, 80))

hr(doc)

# ── Abstract ──────────────────────────────────────────────────────────────────
add_heading(doc, "Abstract / 摘要", level=1)

add_bilingual(doc,
    "We develop a probabilistic surrogate framework for coupled thermo-mechanical "
    "analysis of a heat-pipe-cooled reactor (HPR), enabling end-to-end "
    "uncertainty-to-risk quantification at a fraction of the cost of high-fidelity simulation.",
    "本文针对热管冷却反应堆（HPR）的耦合热–力分析开发了一套概率代理建模框架，可在远低于"
    "高保真仿真计算成本的条件下，实现从材料不确定性到结构风险的端到端定量分析。")

add_bilingual(doc,
    "The surrogate is a heteroscedastic multilayer perceptron (HeteroMLP) trained on "
    "approximately 2900 coupled OpenMC–FEniCS simulations, predicting distributions over "
    "15 physics outputs from 8 material uncertainty inputs across two fidelity levels: "
    "a decoupled first-pass prediction and a converged coupled steady-state.",
    "代理模型为异方差多层感知机（HeteroMLP），基于约2900次耦合OpenMC–FEniCS仿真数据训练，"
    "从8个材料不确定性输入出发，预测两个精度层次（解耦第一轮预测和耦合稳态）上15个物理输出量"
    "的概率分布。")

add_bilingual(doc,
    "A physics-regularized variant embeds Spearman-rank monotonicity constraints and "
    "achieves a test negative log-likelihood of 0.305, with coupled steady-state maximum "
    "global stress R\u00b2 = 0.929 (RMSE = 7.9 MPa).",
    "含物理约束的正则化变体通过嵌入Spearman秩单调性约束，在测试集上达到负对数似然0.305，"
    "耦合稳态最大全局应力拟合优度R\u00b2=0.929（均方根误差7.9 MPa）。")

add_bilingual(doc,
    "Forward propagation of 20,000 prior Monte Carlo samples reveals that multi-physics "
    "coupling substantially compresses uncertainty: the coupled stress mean is 39 MPa lower "
    "than the decoupled prediction (153 MPa vs. 193 MPa), with standard deviation reduced "
    "by 47%; keff uncertainty collapses by 183\u00d7. The predicted probability of exceeding "
    "the 131 MPa design threshold is 0.847 under the prior.",
    "对2万个先验蒙特卡洛样本的前向传播揭示：多物理耦合显著压缩输出不确定性——耦合稳态应力均值"
    "比解耦预测低39 MPa（153 MPa vs. 193 MPa），标准差降低47%；keff不确定性压缩183倍。"
    "在先验认知下，超过131 MPa设计阈值的概率估计为0.847。")

add_bilingual(doc,
    "Sobol sensitivity analysis identifies Young's modulus intercept (E_intercept) as the "
    "dominant stress driver (S\u2081 = 0.598, 90% CI: 0.596\u20130.601) and thermal expansion "
    "coefficient (\u03b1_base) as the dominant keff driver (S\u2081 = 0.775, CI: 0.774\u20130.776), "
    "indicating distinct dominant physical pathways with limited cross-interaction.",
    "Sobol全局敏感性分析识别弹性模量截距（E_intercept）为应力主控因子（S\u2081=0.598，"
    "CI：0.596\u20130.601），热膨胀系数（\u03b1_base）为keff主控因子（S\u2081=0.775，"
    "CI：0.774\u20130.776），表明两个输出量由两条主控因子不同、低阶交叉作用有限的独特物理"
    "传播路径支配。")

add_bilingual(doc,
    "Posterior calibration over 20 benchmark cases demonstrates that, for all ten "
    "extreme-stress scenarios with observed stress \u2265 220 MPa, the posterior probability "
    "of exceeding 131 MPa converges to 1.0. The surrogate achieves a single-sample GPU "
    "latency of 1.7\u00d710\u207b\u2077 s, representing a speedup of more than 23\u00d710\u2076 "
    "over the coupled solver.",
    "对20个基准测试案例的后验标定表明，全部10个真实应力\u2265220 MPa的极端应力场景超过131 MPa"
    "的后验概率均收敛至1.0。代理模型单次GPU推断延迟仅1.7\u00d710\u207b\u2077秒，"
    "比耦合高保真求解器快逾2300万倍。")

hr(doc)

# ── Introduction ──────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction / 引言", level=1)

add_bilingual(doc,
    "Heat-pipe-cooled reactors (HPRs) are an emerging class of compact microreactors in "
    "which the core comprises tightly coupled heat pipes, fuel rods, and monolith structures.",
    "热管冷却反应堆（HPR）是一类新兴紧凑型微型反应堆，其堆芯由热管、燃料棒和单石结构紧密耦合"
    "构成，独特的结构形式使其在材料不确定性传播上具有复杂的多物理场耦合特征。")

add_bilingual(doc,
    "Uncertainties in material properties—arising from manufacturing tolerances and "
    "irradiation-driven degradation—propagate through a coupled neutronics–thermal–structural "
    "feedback loop and can amplify into substantial structural stress risk.",
    "来自制造公差和辐照退化的材料性能不确定性，通过中子物理–热学–结构力学的耦合反馈环传播，"
    "并可能放大为显著的结构应力风险，对反应堆长期安全运行构成威胁。")

add_bilingual(doc,
    "Prior high-fidelity analysis of the nominal MegaPower design reports peak monolith "
    "stresses reaching 169 MPa under normal operating conditions, which exceeds the yield "
    "strength of the SS316 stainless-steel monolith at reactor operating temperatures (131 MPa); "
    "characterising how material-parameter uncertainty translates into exceedance probability "
    "is therefore a primary safety concern.",
    "已有高保真分析表明，标称MegaPower设计在正常工况下的单石结构峰值应力可达169 MPa，超过"
    "SS316不锈钢在工作温度下的屈服强度（131 MPa）；因此，定量表征材料参数不确定性如何转化为"
    "超阈概率是首要安全关切，也是本文工作的直接动机。")

add_bilingual(doc,
    "Quantifying this risk requires sampling over the full 8-dimensional material parameter "
    "space, which is infeasible with repeated high-fidelity (HF) solver evaluations: each "
    "coupled OpenMC–FEniCS run takes approximately one hour.",
    "对完整8维材料参数空间的风险定量需要大量重复的高保真求解器评估，而每次耦合OpenMC–FEniCS"
    "运行约需1小时，使暴力蒙特卡洛方法在实际工程中不具可行性。")

add_bilingual(doc,
    "Existing surrogate approaches face complementary limitations. Gaussian process (GP) "
    "regression and polynomial chaos expansion (PCE) become computationally prohibitive as "
    "the input dimension grows, and require additional regularity assumptions that may not hold "
    "across the full uncertainty range. Physics-informed neural networks (PINNs) are designed "
    "for problems where exact governing equations can be expressed as residuals; the iterative "
    "coupling boundary conditions in an HPR preclude this formulation. Data-driven neural "
    "surrogates offer scalable inference but typically lack uncertainty calibration and "
    "physical consistency guarantees.",
    "现有代理方法各有局限：高斯过程（GP）回归和多项式混沌展开（PCE）在输入维度增大时计算代价"
    "迅速上升，且依赖在全不确定性范围内可能不成立的正则性假设；物理信息神经网络（PINN）适用于"
    "支配方程可表达为残差形式的问题，HPR中的迭代耦合边界条件不满足这一前提；纯数据驱动神经代理"
    "虽可扩展推断，但通常缺乏不确定性校准和物理一致性保证。")

add_bilingual(doc,
    "This work makes four contributions. (1) We develop a HeteroMLP surrogate that jointly "
    "predicts predictive mean and variance for each of 15 physics outputs, enabling calibrated "
    "probabilistic inference at the cost of a single forward pass. (2) We propose a "
    "physics-regularized training objective augmenting the Gaussian NLL with a Spearman-rank "
    "monotonicity term, demonstrating substantially improved OOD robustness for stress under "
    "\u03b1_base extrapolation (\u0394R\u00b2 = +0.170). (3) We present a complete "
    "uncertainty-to-risk pipeline: forward Monte Carlo propagation, global Sobol sensitivity "
    "analysis with bootstrap confidence intervals, and MCMC-based posterior calibration. "
    "(4) We quantify—for the first time in this system—the degree to which multi-physics "
    "coupling compresses material-parameter-driven uncertainty, revealing a 183\u00d7 reduction "
    "in keff spread and a 47% reduction in stress standard deviation.",
    "本文作出四项贡献：（1）开发HeteroMLP代理模型，对15个物理输出量同时预测均值和方差，"
    "以单次前向传播实现经校准的概率推断；（2）提出引入Spearman秩单调性惩罚项的物理正则化训练"
    "目标，在\u03b1_base外推下取得\u0394R\u00b2=+0.170的分布外精度提升；（3）构建完整的不确定性"
    "–风险分析流水线，包含前向蒙特卡洛传播、带Bootstrap置信区间的全局Sobol敏感性分析和MCMC"
    "后验标定；（4）定量揭示多物理耦合对材料参数驱动不确定性的压缩效应——keff标准差压缩"
    "183倍，应力标准差降低47%。")

hr(doc)

# ── RESULTS (NCS: Results before Methods) ─────────────────────────────────────
add_heading(doc, "2. Results / 结果", level=1)

# 2.1
add_heading(doc, "2.1  Surrogate Accuracy and Physics Regularization / 代理精度与物理正则化效果", level=2)

add_bilingual(doc,
    "The physics-regularized surrogate achieves test NLL = 0.305 and, for the primary "
    "safety metric (coupled steady-state maximum global stress), R\u00b2 = 0.929, "
    "RMSE = 7.9 MPa, and PICP\u2089\u2080 = 0.913; mean R\u00b2 across the five primary "
    "outputs is 0.795. This model is selected as the primary surrogate for all subsequent "
    "analyses on the basis of its distributional calibration and robustness under "
    "distributional shift, demonstrated below.",
    "物理正则化代理测试集NLL=0.305；主要安全指标（耦合稳态最大全局应力）R\u00b2=0.929，"
    "RMSE=7.9 MPa，PICP\u2089\u2080=0.913；5个主要输出量平均R\u00b2为0.795。"
    "以下验证表明，该模型在分布内标定精度和分布外鲁棒性上均优于基线，因此作为后续所有分析"
    "的主代理模型。")

add_bilingual(doc,
    "Temperature outputs attain R\u00b2 \u2248 0.58\u20130.59, reflecting the inherently "
    "lower sensitivity of converged temperatures to the 8 material inputs in the coupled "
    "regime; absolute RMSE remains below 5 K. Wall expansion achieves R\u00b2 = 0.996. "
    "For keff, the surrogate reaches R\u00b2 = 0.874.",
    "温度输出R\u00b2约0.58\u20130.59，反映了耦合稳态下温度对8个材料参数的固有低敏感性；"
    "绝对RMSE仍低于5 K。壁面膨胀量R\u00b2=0.996；keff的R\u00b2=0.874。")

add_bilingual(doc,
    "For \u03b1_base OOD extrapolation, the physics-regularized model retains coupled stress "
    "R\u00b2 = 0.944 vs. 0.774 for the baseline, a difference of \u0394R\u00b2 = +0.170 "
    "directly attributable to the enforced \u03b1_base\u2013stress monotonicity constraint. "
    "This improvement is specific to the stress response under \u03b1_base extrapolation; "
    "OOD performance on other outputs and input features is heterogeneous (see Appendix).",
    "\u03b1_base外推场景下，物理正则化模型保持耦合应力R\u00b2=0.944，基线降至0.774"
    "（\u0394R\u00b2=+0.170），此提升直接源于所强制的\u03b1_base\u2192应力单调性约束。"
    "该提升专指\u03b1_base外推下的应力响应；其他输出量和输入特征的分布外性能改善情况"
    "不尽相同（见附录）。")

add_table_placeholder(doc, "Table 1",
    "Per-output accuracy of the physics-regularized surrogate (monotone-regularized surrogate) on the held-out "
    "test set (n = 435). Primary outputs marked (\u2605).",
    "单调正则化代理测试集各输出精度：MAE / RMSE / R\u00b2 / PICP\u2089\u2080 / "
    "MPIW\u2089\u2080。主要输出量以★标注。")

add_table_placeholder(doc, "Table 2",
    "Out-of-distribution evaluation: coupled steady-state stress R\u00b2 for held-out "
    "samples with extreme values of each input feature.",
    "分布外评估：各输入特征极端值下的耦合稳态应力R\u00b2（data-only baseline vs. monotone-regularized surrogate）。")

add_fig_placeholder(doc, "Figure 1 (SVG)",
    "Surrogate accuracy: (A) parity plot for coupled stress (R\u00b2=0.929, RMSE=7.9 MPa); "
    "(B) R\u00b2 for all five primary outputs; (C) OOD robustness for \u03b1_base extrapolation.",
    "代理精度三面板：(A) 应力预测–真值散点图；(B) 5个主要输出R\u00b2对比；"
    "(C) \u03b1_base外推下OOD鲁棒性。")

# 2.2
add_heading(doc, "2.2  Forward Uncertainty Propagation and Threshold Risk / 前向不确定性传播与阈值风险", level=2)

add_bilingual(doc,
    "In the decoupled prediction, maximum global stress: \u03bc = 192.7 MPa, "
    "\u03c3 = 40.9 MPa (q\u2095\u2085 = 264.8 MPa); keff standard deviation = 0.0625 "
    "(\u2248 6250 pcm).",
    "解耦预测：最大全局应力\u03bc=192.7 MPa，\u03c3=40.9 MPa（q\u2095\u2085=264.8 MPa）；"
    "keff标准差0.0625（约6250 pcm）。")

add_bilingual(doc,
    "In the coupled steady-state, the stress mean falls to 153.4 MPa (\u221239.3 MPa, "
    "\u221220%) and the standard deviation decreases from 40.9 to 21.7 MPa (47% reduction). "
    "keff standard deviation is 3.42\u00d710\u207b\u2074 (\u2248 34 pcm), approximately "
    "183\u00d7 smaller than the decoupled value.",
    "耦合稳态：应力均值降至153.4 MPa（\u221239.3 MPa，降幅20%），标准差降至21.7 MPa"
    "（降幅47%）。keff标准差为3.42\u00d710\u207b\u2074（约34 pcm），约为解耦值的1/183。")

add_bilingual(doc,
    "The mechanism is a neutronics\u2013thermal negative feedback: high \u03b1_base "
    "expands the core, reduces neutron leakage, raises keff, and moderates local power "
    "density, partially offsetting the stress increase. This feedback is absent in the "
    "decoupled prediction; omitting it inflates stress standard deviation by approximately 2\u00d7.",
    "机制为中子–热负反馈：高\u03b1_base引起堆芯膨胀、减少中子泄漏、提升keff并调节局部功率"
    "密度，部分抵消热应力增大。解耦预测中不存在这一反馈，忽略它将使应力标准差被高估约两倍。")

add_bilingual(doc,
    "P_exceed(131 MPa) = 0.847 under the prior (epistemic, mean predictions). This matches "
    "prior HF analysis reporting peak monolith stress of 169 MPa for the nominal design; "
    "the training dataset predominantly samples the above-threshold regime, and "
    "P_exceed = 0.847 reflects the position of the nominal design relative to the "
    "material yield limit, not a surrogate artefact.",
    "先验认知超阈概率P_exceed(131 MPa)=0.847（仅使用预测均值）。这与标称设计峰值应力"
    "169 MPa的已有高保真分析一致；训练数据集主要采样阈值以上的运行区间，P_exceed=0.847"
    "反映标称设计点相对于材料屈服极限的位置，而非代理伪象。")

add_table_placeholder(doc, "Table 3",
    "Decoupled vs. coupled steady-state uncertainty distributions for key outputs, "
    "from 20,000 prior Monte Carlo samples (monotone-regularized surrogate surrogate, epistemic mean predictions).",
    "解耦与耦合稳态不确定性分布对比（均值、标准差、q\u2095\u2085）——20000个先验MC样本。")

add_fig_placeholder(doc, "Figure 2 (SVG)",
    "Forward uncertainty propagation: (A) stress histograms decoupled vs. coupled; "
    "(B) keff decoupled distribution; (C) keff coupled distribution (183\u00d7 zoomed).",
    "前向UQ三面板：(A) 应力直方图解耦vs耦合；(B) keff解耦分布；(C) keff耦合放大183\u00d7。")

# 2.3
add_heading(doc, "2.3  Global Sensitivity Analysis / 全局敏感性分析", level=2)

add_bilingual(doc,
    "For coupled steady-state maximum global stress (monotone-regularized surrogate, data-mono): E_intercept is the "
    "dominant driver (S\u2081 = 0.598, 90% CI: 0.596\u20130.601), accounting for \u224860% of "
    "first-order variance. \u03b1_base ranks second (S\u2081 = 0.148, CI: 0.144\u20130.151), "
    "and k_ref,SS316 is third (S\u2081 = 0.080, CI: 0.077\u20130.083). The remaining five "
    "inputs have confidence intervals that include or approach zero.",
    "耦合稳态最大全局应力（monotone-regularized surrogate，data-mono）：E_intercept最主导（S\u2081=0.598，"
    "CI：0.596\u20130.601），约占一阶方差的60%；\u03b1_base居第二（S\u2081=0.148，"
    "CI：0.144\u20130.151）；k_ref,SS316第三（S\u2081=0.080，CI：0.077\u20130.083）；"
    "其余5个参数的置信区间均包含或接近零。")

add_bilingual(doc,
    "The dominance of E_intercept reflects its role as the primary stiffness parameter: "
    "a higher Young's modulus intercept directly amplifies thermal stress under the same "
    "temperature gradient, independently of the coupling pathway. The non-zero contribution "
    "of k_ref,SS316 (S\u2081 = 0.080) reflects a real thermal-conduction pathway partially "
    "masked by the stronger elastic effect.",
    "E_intercept的主导反映其作为主要刚度参数的作用：弹性模量截距越高，在相同温度梯度下"
    "直接放大热应力，与耦合路径无关。k_ref,SS316（S\u2081=0.080）的非零贡献反映真实的"
    "热传导路径，但被更强的弹性效应部分掩盖。")

add_bilingual(doc,
    "For coupled steady-state keff (monotone-regularized surrogate): \u03b1_base is the single dominant factor "
    "(S\u2081 = 0.775, CI: 0.774\u20130.776); \u03b1_slope is secondary (S\u2081 = 0.177, "
    "CI: 0.174\u20130.180). The remaining six parameters have confidence intervals including zero "
    "and are not treated as stable contributors. The dominance of \u03b1_base reflects that "
    "thermal expansion governs core geometry and thereby neutron leakage, the principal "
    "driver of keff variation.",
    "耦合稳态keff（monotone-regularized surrogate）：\u03b1_base唯一主控（S\u2081=0.775，CI：0.774\u20130.776）；"
    "\u03b1_slope次之（S\u2081=0.177，CI：0.174\u20130.180）；其余6个参数的置信区间均包含零，"
    "不作为稳定贡献因子。\u03b1_base的主导反映热膨胀控制堆芯几何形变进而决定中子泄漏率，"
    "是keff变化的主要驱动因素。")

add_bilingual(doc,
    "The two primary outputs are governed by distinct physical pathways with limited "
    "cross-interaction at the dominant-factor level: elastic stiffness (E_intercept) governs "
    "stress, while thermal expansion (\u03b1_base) governs reactivity. Design interventions "
    "targeting stress (e.g., tighter E_intercept manufacturing tolerance) therefore do not "
    "substantially perturb keff, and vice versa, enabling partially decoupled design strategies.",
    "两个主要输出量由不同物理路径支配，在主控因子层面交叉作用有限：弹性刚度（E_intercept）"
    "主控应力，热膨胀（\u03b1_base）主控反应性。因此，针对应力（如收紧E_intercept制造公差）"
    "的设计干预不会显著扰动keff行为，反之亦然，两条路径可采用基本分离的设计管控策略。")

add_table_placeholder(doc, "Table 4",
    "First-order Sobol indices S\u2081 (90% CI, data-mono monotone-regularized surrogate) — only factors with CI not "
    "including zero listed; full results in Appendix. "
    "Stress: E_intercept 0.598 (0.596\u20130.601), \u03b1_base 0.148, k_ref,SS316 0.080. "
    "keff: \u03b1_base 0.775 (0.774\u20130.776), \u03b1_slope 0.177.",
    "Sobol一阶指数S\u2081（90% CI，data-mono monotone-regularized surrogate）——仅列CI不包含零的因子。"
    "应力：E_intercept 0.598，\u03b1_base 0.148，k_ref,SS316 0.080。"
    "keff：\u03b1_base 0.775，\u03b1_slope 0.177。")

add_fig_placeholder(doc, "Figure 3 (SVG)",
    "Global Sobol first-order sensitivity: (A) max global stress — E_intercept dominant "
    "(S\u2081=0.598); (B) keff — \u03b1_base dominant (S\u2081=0.775). Grey bars: CI includes zero.",
    "全局Sobol敏感性分析：(A) 应力一阶指数——E_intercept主控（S\u2081=0.598）；"
    "(B) keff一阶指数——\u03b1_base主控（S\u2081=0.775）。灰色柱：置信区间含零。")

# 2.4
add_heading(doc, "2.4  Observation-Driven Posterior Calibration / 观测驱动后验标定", level=2)

add_bilingual(doc,
    "MCMC calibration was performed on 20 benchmark test cases (observed stress "
    "\u2248100\u2013226 MPa) and 10 extreme-stress cases (observed stress \u2265220 MPa). "
    "For each case, the Metropolis-Hastings sampler ran for 8000 total iterations with "
    "2000 burn-in steps discarded and thinning factor 5, yielding 1200 effective posterior "
    "samples. Acceptance rates ranged from 0.47 to 0.61, consistent with adequate single-chain "
    "Metropolis-Hastings exploration for the 4-parameter posterior. Convergence was assessed by "
    "inspecting trace plots: all four calibrated parameters exhibit stationary random "
    "fluctuations without directional drift after the burn-in period.",
    "MCMC标定在20个基准测试案例（观测应力约100\u2013226 MPa）和10个极端应力案例"
    "（观测应力\u2265220 MPa）上执行。每个案例运行8000步迭代，丢弃前2000步烧入期，"
    "每5步保留1个样本，共获得1200个有效后验样本。接受率0.47\u20130.61，与4参数后验的"
    "MH充分探索要求相符。通过trace plot检验收敛性：烧入期结束后，4个标定参数均呈现"
    "平稳随机波动，无方向性漂移，确认链已充分探索后验分布。")

add_bilingual(doc,
    "Note: posterior predictions are validated against the nearest-neighbour entry in the "
    "HF dataset rather than a true HF rerun at the posterior mean (proxy validation).",
    "注：后验预测与观测的比较采用最近邻HF数据集查找方式（proxy验证），而非在后验均值处"
    "重新运行高保真求解器，本节结果均在此前提下解读。")

add_bilingual(doc,
    "The marginal posterior distributions of the four calibrated parameters exhibit distinct "
    "non-Gaussian features and parameter compensation effects. For high-stress cases, "
    "E_intercept is skewed toward lower values\u2014reduced stiffness provides stress "
    "relief\u2014while \u03b1_base concentrates at the lower end, as lower thermal expansion "
    "suppresses the coupling-driven stress amplification. The joint posterior of E_intercept "
    "and \u03b1_base shows a pronounced compensation effect: feasible solutions cluster where "
    "both parameters are simultaneously small, matching the multiplicative stress dependence "
    "\u03c3 \u221d E \u00b7 \u03b1 \u00b7 T.",
    "4个标定参数的边际后验分布表现出明显的非高斯特征和参数补偿效应。高应力案例中，E_intercept"
    "偏向低值（较低刚度提供应力释放），\u03b1_base集中在低端（较小热膨胀抑制耦合驱动的应力放大）；"
    "E_intercept与\u03b1_base的联合后验呈强补偿效应——可行解集中于两者同时偏小的区域，"
    "与乘积形式\u03c3\u221dE\u00b7\u03b1\u00b7T吻合。")

add_bilingual(doc,
    "P_safe(131 MPa) decreases systematically with increasing observed stress: "
    "high-stress cases \u2192 P_safe \u2248 0; low-stress cases \u2192 P_safe \u2248 1.0.",
    "P_safe(131 MPa)随观测应力升高系统性降低：高应力观测对应P_safe\u22480，"
    "低应力观测对应P_safe\u22481.0，系统能够从单次应力观测中正确推断安全风险等级。")

add_bilingual(doc,
    "For all 10 extreme-stress cases (true stress 220.5\u2013278.4 MPa), the posterior "
    "exceeding probability converges to P_exceed = 1.0, compared to prior range 0.841\u20130.957. "
    "The surrogate prediction at the posterior mean closely tracks the true stress "
    "(MAE < 6 MPa), confirming that the calibration localises the posterior in the correct "
    "high-risk parameter region. Sobol analysis identifies E_intercept as the dominant "
    "stress driver (S\u2081 = 0.598); following extreme-stress observations, the posterior "
    "systematically shifts E_intercept and \u03b1_base into the high-stress parameter region, "
    "consistent with the multiplicative dependence \u03c3 \u221d E\u00b7\u03b1\u00b7T.",
    "全部10个极端应力案例（真实应力220.5\u2013278.4 MPa）的超阈后验概率均收敛至P_exceed=1.0，"
    "而先验概率范围仅为0.841\u20130.957。后验均值处的代理预测与真实应力高度吻合（MAE<6 MPa），"
    "确认标定将后验定位于正确的高风险参数区域。Sobol分析识别E_intercept为应力主控因子"
    "（S\u2081=0.598），极端应力观测后，后验系统性地将E_intercept和\u03b1_base推入高应力"
    "参数区域，与乘积形式\u03c3\u221dE\u00b7\u03b1\u00b7T相符。")

add_table_placeholder(doc, "Table 5",
    "Extreme-stress posterior risk updating: prior and posterior probabilities of exceeding "
    "131 MPa for 10 benchmark cases with true stress \u2265 220 MPa.",
    "极端应力后验风险更新——10个案例的真实应力、P_exceed\u207d\u1d56\u02b3\u1d52\u02b3、"
    "P_exceed\u207d\u1d56\u1d52\u02e2\u1d57。")

add_fig_placeholder(doc, "Figure 4 (SVG)",
    "Posterior calibration: (A) posterior mean stress prediction vs. observed stress "
    "(nearest-neighbour HF proxy, 20 benchmark cases); (B) prior and posterior exceedance "
    "probabilities P_exceed(131 MPa) for the 10 extreme-stress cases.",
    "后验标定结果：(A) 后验均值预测vs观测应力散点（最近邻proxy验证）；"
    "(B) 10个极端应力案例先验/后验P_exceed对比。")

hr(doc)

# ── Discussion ────────────────────────────────────────────────────────────────
add_heading(doc, "3. Discussion / 讨论", level=1)

add_bilingual(doc,
    "Multi-physics coupling as an uncertainty modulator. The 47% reduction in stress "
    "standard deviation from decoupled to coupled steady-state reflects a genuine negative "
    "feedback mechanism in the coupled physics, not a model artefact. Performing UQ at the "
    "decoupled level would overestimate stress standard deviation by approximately a factor "
    "of two, leading to inflated tail-risk estimates. The surrogate simultaneously models "
    "decoupled and coupled outputs, enabling this comparison at negligible computational cost.",
    "多物理耦合作为不确定性调节机制。47%的应力标准差降低是真实的物理负反馈效应，非模型伪象。"
    "在解耦层次开展不确定性传播将系统性高估应力标准差约两倍，导致尾部风险被夸大。"
    "代理模型同时建模解耦和耦合输出，使这一比较几乎不产生额外计算开销。")

add_bilingual(doc,
    "Physics regularization and model selection. Physics regularization is preferred on "
    "three grounds. First, in-distribution accuracy is comparable to the baseline "
    "(\u0394R\u00b2 \u2248 +0.005; Appendix). Second, under \u03b1_base extrapolation the "
    "regularized model retains R\u00b2 = 0.944 while the baseline falls to 0.774 "
    "(\u0394R\u00b2 = +0.170); the unconstrained baseline violates the physical monotonicity "
    "in low-density input regions, and the Spearman-rank regularizer corrects this directly. "
    "Third, the Sobol first-order index for E_intercept shifts from S\u2081 = 0.613 "
    "(baseline, data-only baseline) to S\u2081 = 0.598 (regularized, monotone-regularized surrogate), with attribution "
    "redistributed to physically consistent secondary factors (\u03b1_base, k_ref,SS316). "
    "Both Sobol results are in the Appendix.",
    "物理正则化与模型选择。物理正则化从三个方面具有优势：第一，分布内精度与基线相当"
    "（\u0394R\u00b2\u22480.005，见附录）；第二，\u03b1_base外推下正则化模型保持R\u00b2=0.944，"
    "基线降至0.774（\u0394R\u00b2=+0.170），Spearman秩正则化直接修正基线在低密度区域的"
    "物理单调性违反；第三，E_intercept的一阶Sobol指数从基线S\u2081=0.613降至正则化S\u2081=0.598，"
    "归因重新分配至物理一致的次级因子（\u03b1_base、k_ref,SS316），两套Sobol结果均见附录。")

add_bilingual(doc,
    "Computational efficiency. The surrogate achieves a single-sample CPU latency of "
    "\u22481.55\u00d710\u207b\u2074 s and a per-sample GPU batch latency of "
    "\u22481.70\u00d710\u207b\u2077 s, compared to \u22483600 s for the coupled HF solver. "
    "The resulting speedup is 23.2\u00d710\u2076 (single CPU) or 21.2\u00d710\u2079 "
    "(GPU batched), enabling full UQ analyses within minutes of wall-clock time. "
    "[Note: hardware specifications to be added.]",
    "计算效率。单次CPU约1.55\u00d710\u207b\u2074秒，GPU批推断约1.70\u00d710\u207b\u2077秒，"
    "对比高保真求解器约2000秒（约33分钟）。加速比23.2\u00d710\u2076（CPU）或21.2\u00d710\u2079"
    "（GPU批量），可在数分钟内完成完整UQ分析。【待补充：具体硬件规格】")

add_bilingual(doc,
    "Implications for design and operation. E_intercept accounts for approximately 60% of "
    "coupled stress variance (S\u2081 = 0.598); tighter manufacturing control of Young's "
    "modulus intercept is therefore the primary lever for stress risk reduction. k_ref,SS316 "
    "contributes a secondary but non-negligible share (S\u2081 = 0.080). \u03b1_base dominates "
    "keff (S\u2081 = 0.775) but contributes only moderately to stress (S\u2081 = 0.148): the "
    "two physical pathways support partially decoupled measurement and control strategies, "
    "reducing cross-interference in the safety-analysis workflow. Posterior calibration "
    "demonstrates that a single in-situ stress observation shifts P_exceed from 0.84\u20130.96 "
    "to 1.0 for extreme-stress cases, supporting sequential experimental design.",
    "设计与运行含义。E_intercept（弹性模量截距）占耦合应力方差约60%（S\u2081=0.598），"
    "是应力风险降低的首要杠杆，因此弹性模量的制造公差管控应被优先收紧。k_ref,SS316贡献"
    "次要但不可忽视的份额（S\u2081=0.080）。\u03b1_base主控keff（S\u2081=0.775），但对应力"
    "贡献适中（S\u2081=0.148），两条物理路径支持基本分离的测量与管控策略，减少安全分析中"
    "的交叉干扰。后验标定表明，单次原位应力观测可将极端应力案例的P_exceed从0.84\u20130.96"
    "提升至1.0，支持序贯实验设计。")

add_bilingual(doc,
    "Limitations. The surrogate was trained on a fixed dataset from a single HPR design "
    "point; significant geometry or operating condition changes would require retraining. "
    "The OOD tests show that monotone-regularized surrogate degrades more gracefully than data-only baseline in extrapolation, "
    "but absolute accuracy still decreases outside the training range. The MCMC calibration "
    "fixes four of the eight input parameters at their true test-case values (SS316 thermal "
    "properties), treating them as known from material characterisation; in practice these "
    "values would not be precisely known, so the reported posterior uncertainty is a lower bound. "
    "A full 8-dimensional calibration treating k_ref,SS316 as unknown would increase posterior "
    "uncertainty but require independent thermal conductivity measurements. The posterior "
    "validation uses the existing HF dataset as a nearest-neighbour proxy rather than a true "
    "HF rerun at the posterior mean.",
    "局限性。代理模型在单一HPR设计点的固定数据集上训练，重大几何或工况变化需重新训练。"
    "OOD测试表明monotone-regularized surrogate在外推时退化更为平缓，但训练范围外绝对精度仍下降。MCMC标定将4个"
    "SS316热性质参数（E_slope, T_ref,SS316, k_ref,SS316, k_slope,SS316）固定在真实测试案例"
    "输入值，等效于假设这些量可由材料表征预先知晓；实际工程中这些值未必精确已知，因此报告的"
    "后验不确定性为下界。若将k_ref,SS316也纳入标定需要独立的热导率测量，将增大后验不确定性。"
    "后验验证采用最近邻HF数据集查找，而非在后验均值处重新运行高保真求解器。")

hr(doc)

# ── Conclusion ────────────────────────────────────────────────────────────────
add_heading(doc, "4. Conclusion / 结论", level=1)

add_bilingual(doc,
    "This work presents a probabilistic surrogate framework for coupled multi-physics "
    "uncertainty quantification in a heat-pipe-cooled reactor, demonstrating four "
    "principal findings:",
    "本文提出了热管冷却反应堆耦合多物理场不确定性量化的概率代理框架，揭示了四项主要发现：")

conclusions_en = [
    "Physics regularization improves generalization. The monotone-regularized surrogate HeteroMLP (Spearman-rank "
    "monotonicity constraints) matches in-distribution accuracy (test NLL 0.305, stress "
    "R\u00b2 = 0.929) while substantially improving OOD performance for \u03b1_base "
    "extrapolation (R\u00b2: 0.774 \u2192 0.944). This supports embedding domain knowledge "
    "as a soft regularizer.",
    "Multi-physics coupling compresses uncertainty. Forward propagation shows coupled "
    "steady-state outputs have significantly tighter distributions: 47% stress standard "
    "deviation reduction and 183\u00d7 keff uncertainty compression. These are consequences "
    "of the neutronics\u2013thermal feedback mechanism and cannot be captured by "
    "single-physics surrogates.",
    "Distinct sensitivity pathways exist. Sobol analysis identifies E_intercept as the "
    "dominant stress driver (S\u2081 = 0.598) and \u03b1_base as the dominant keff driver "
    "(S\u2081 = 0.775), indicating that design interventions for stress reduction and "
    "reactivity control can be treated with largely separable strategies at the "
    "dominant-factor level.",
    "Posterior calibration sharpens risk assessments. MCMC-based parameter updating "
    "conditioned on observations converges P_exceed(131 MPa) to 1.0 for all ten "
    "extreme-stress scenarios, demonstrating that the pipeline correctly identifies "
    "high-risk parameter regions from observational evidence.",
]
conclusions_cn = [
    "物理正则化改善泛化性。monotone-regularized surrogate HeteroMLP（Spearman秩单调性约束）分布内精度与基线相当"
    "（NLL 0.305，应力R\u00b2=0.929），同时大幅提升\u03b1_base外推的分布外性能"
    "（R\u00b2：0.774\u21920.944）。",
    "多物理耦合压缩不确定性。应力标准差降低47%，keff不确定性压缩183倍——中子–热负反馈机制"
    "的直接体现，单物理代理无法捕捉。",
    "存在独特的敏感性传播路径。E_intercept主控应力（S\u2081=0.598），\u03b1_base主控keff"
    "（S\u2081=0.775），在主控因子层面两条路径交叉作用有限，应力控制与反应性控制的设计干预"
    "措施可采用基本分离的策略。",
    "后验标定精化风险评估。MCMC参数更新使全部10个极端应力场景的P_exceed(131 MPa)收敛至1.0，"
    "证明流水线能够从观测证据中正确识别高风险参数区域。",
]
for i, (en, cn) in enumerate(zip(conclusions_en, conclusions_cn), 1):
    add_bilingual(doc, f"({i}) {en}", f"（{i}）{cn}")

add_bilingual(doc,
    "The computational speedup (> 23\u00d710\u2076) makes this framework practical for "
    "iterative design exploration and real-time risk monitoring in reactor systems with "
    "significant material uncertainty.",
    "2300万倍以上的计算加速比使该框架适用于具有显著材料不确定性的反应堆系统的迭代设计探索"
    "和实时风险监控，为核能系统的概率安全评价提供了一条兼具物理可信性和计算可行性的新途径。")

hr(doc)

# ── Methods (NCS: after Results) ──────────────────────────────────────────────
add_heading(doc, "5. Methods / 方法", level=1)

# 5.1
add_heading(doc, "5.1  Material Uncertainty Model / 材料不确定性模型", level=2)

add_bilingual(doc,
    "Eight material parameters are treated as independent uncertain inputs: the "
    "temperature-dependent Young's modulus E(T) = E_slope\u00b7T + E_intercept "
    "(slope and intercept separately), Poisson's ratio \u03bd, thermal expansion "
    "coefficient \u03b1(T) = \u03b1_base + \u03b1_slope\u00b7T (base and slope separately), "
    "and three SS316 stainless-steel thermal properties (T_ref,SS316, k_ref,SS316, "
    "k_slope,SS316). Here k_slope,SS316 is the temperature coefficient of SS316 thermal "
    "conductivity\u2014the slope in k_SS316(T) = k_ref,SS316 + k_slope,SS316\u00b7(T \u2212 "
    "T_ref,SS316)\u2014and is distinct from the structural thermal expansion parameters "
    "\u03b1_base and \u03b1_slope. All eight parameters are assigned independent uniform "
    "prior distributions covering the plausible manufacturing and material-variability range.",
    "8个材料参数被视为独立不确定性输入：温度相关弹性模量E(T)=E_slope\u00b7T+E_intercept"
    "（斜率和截距分别处理）、泊松比\u03bd、热膨胀系数\u03b1(T)=\u03b1_base+\u03b1_slope\u00b7T"
    "（基值和斜率分别处理）、以及SS316不锈钢的三个热性质（T_ref,SS316, k_ref,SS316, "
    "k_slope,SS316）。其中k_slope,SS316为SS316热导率的温度系数，与结构热膨胀系数参数\u03b1_base"
    "和\u03b1_slope物理含义完全不同。全部8个参数均被赋予独立均匀先验分布，覆盖制造公差"
    "和材料变异性的合理范围。")

add_fig_placeholder(doc, "Figure M1 (placeholder)",
    "HPR reactor cross-section showing monolith structure, fuel rods, heat-pipe "
    "locations, and SS316 cladding. Annotate material regions corresponding to the "
    "8 uncertain input parameters in Table A1.",
    "[占位] HPR反应堆横截面——单石结构、燃料棒、热管位置及SS316包壳。"
    "标注与表A1中8个不确定输入参数对应的材料区域。")

add_fig_placeholder(doc, "Figure M2 (placeholder)",
    "HeteroMLP architecture: 8-dim input \u03b8 \u2192 shared hidden layers \u2192 "
    "twin output heads for \u03bc(\u03b8) and log\u03c3\u00b2(\u03b8) (15 outputs each). "
    "Monotone-regularized variant adds Spearman-rank penalty during training.",
    "[占位] HeteroMLP架构图：8维输入\u03b8 \u2192 共享隐层 \u2192 "
    "均值\u03bc(\u03b8)与对数方差log\u03c3\u00b2(\u03b8)双输出头（各15个输出）。"
    "单调正则化变体在训练时附加Spearman秩惩罚项。")

add_table_placeholder(doc, "Table A1 / 表A1  —  Input parameters and prior distributions",
    "All priors are independent uniform U[0.9\u00d7nominal, 1.1\u00d7nominal] (\u00b110%). "
    "Symbols: E_slope \u22127.0\u00d710\u2077 Pa/K [\u22127.70, \u22126.30]\u00d710\u2077; "
    "E_intercept 2.0\u00d710\u00b9\u00b9 Pa [1.80, 2.20]\u00d710\u00b9\u00b9; "
    "\u03bd 0.31 [0.279, 0.341]; "
    "\u03b1_base 1.0\u00d710\u207b\u2075 K\u207b\u00b9 [9.00, 11.00]\u00d710\u207b\u2076; "
    "\u03b1_slope 5.0\u00d710\u207b\u2079 K\u207b\u00b2 [4.50, 5.50]\u00d710\u207b\u2079; "
    "T_ref,SS316 923.15 K [830.8, 1015.5]; "
    "k_ref,SS316 23.2 W/(m\u00b7K) [20.88, 25.52]; "
    "k_slope,SS316 1/75\u22480.0133 W/(m\u00b7K\u00b2) [0.0120, 0.0147].",
    "所有先验均为独立均匀分布 U[0.9\u00d7\u03b8_nom, 1.1\u00d7\u03b8_nom]（\u00b110%摄动）。"
    "8个参数：E_slope、E_intercept、\u03bd、\u03b1_base、\u03b1_slope、T_ref,SS316、"
    "k_ref,SS316、k_slope,SS316。标称值与先验区间见正文。")

# 5.2
add_heading(doc, "5.2  Coupled High-Fidelity Simulation / 耦合高保真仿真", level=2)

add_bilingual(doc,
    "The high-fidelity workflow couples OpenMC Monte Carlo neutronics with the FEniCS "
    "finite-element solver in an iterative feedback loop. The neutronics calculation yields "
    "keff and the volumetric power distribution, which drives the FEniCS thermal\u2013"
    "structural computation. The resulting temperature fields update material properties and "
    "geometry, which feeds back into neutronics. Iteration continues until keff and the "
    "temperature field converge to a self-consistent steady state. Each full evaluation "
    "requires approximately one hour.",
    "高保真工作流以迭代反馈方式耦合OpenMC蒙特卡洛中子学求解器和FEniCS有限元求解器。中子学"
    "计算给出keff和体积功率分布，驱动FEniCS热–结构计算；所得温度场更新材料性质和几何构型，"
    "反馈回中子学计算，迭代至keff和温度场达到自洽稳态。每次完整评估约需1小时。")

add_bilingual(doc,
    "The coupled HPR system exhibits two distinct physical pathways through which "
    "material-parameter uncertainty propagates to safety-critical outputs. "
    "(i) Stress pathway: k_ref,SS316 governs heat transport from fuel to heat pipe, setting "
    "temperature gradients and thermal stress; this pathway has weak coupling to neutronics "
    "feedback at first order. "
    "(ii) Reactivity pathway: \u03b1_base controls core geometric expansion under heating, "
    "modifying neutron leakage and keff. In the coupled steady-state, high \u03b1_base "
    "expands the core, reduces leakage, raises keff, and increases local power density. "
    "These pathways interact through a negative feedback: high thermal expansion raises "
    "keff and power deposition, partially offsetting the stress increase.",
    "耦合HPR系统中材料不确定性通过两条独特的物理路径传播：（i）应力路径——k_ref,SS316控制"
    "从燃料到热管的热传递效率，决定温度梯度和热应力大小，一阶上与中子物理反馈耦合较弱；"
    "（ii）反应性路径——\u03b1_base控制加热时的堆芯几何膨胀，改变中子泄漏率和keff，在耦合"
    "稳态中高\u03b1_base通过膨胀–泄漏–功率链形成负反馈。两条路径通过负反馈相互作用：高热膨胀"
    "引起keff和功率沉积升高，部分抵消应力增大。")

add_bilingual(doc,
    "A dataset of n = 2900 samples was generated by Latin hypercube sampling of the "
    "8-dimensional input space. Each sample provides 15 physics outputs across two fidelity "
    "levels: decoupled prediction (decoupled single-pass) and coupled steady-state (coupled steady-state). The dataset "
    "was split into training (n_train = 2029), validation (n_val = 436), and test "
    "(n_test = 435) sets using a fixed random seed (2026) prior to any model selection.",
    "通过8维输入空间的拉丁超立方采样生成2900个样本的数据集，每个样本提供15个物理输出量。"
    "数据集在任何模型选择之前以固定随机种子（2026）划分为训练集（2029个）、验证集（436个）"
    "和测试集（435个）。")

# 5.3
add_heading(doc, "5.3  Physics-Regularized Probabilistic Surrogate / 物理正则化概率代理模型", level=2)

add_bilingual(doc,
    "The surrogate is a HeteroMLP that maps \u03b8 \u2208 \u211d\u2078 to predictive mean "
    "\u03bc(\u03b8) and log-variance log \u03c3\u00b2(\u03b8) for each of 15 physics outputs. "
    "At test time: p(y_j|\u03b8) = \u2115(\u03bc_j(\u03b8), \u03c3_j\u00b2(\u03b8)).",
    "代理模型为HeteroMLP，将8维参数向量\u03b8映射为15个物理输出量各自的预测均值\u03bc(\u03b8)"
    "和对数方差log \u03c3\u00b2(\u03b8)。测试时预测分布为p(y_j|\u03b8)=\u2115(\u03bc_j(\u03b8),"
    "\u03c3_j\u00b2(\u03b8))，即各输出量服从以代理预测为参数的高斯分布。")

add_bilingual(doc,
    "The baseline surrogate (data-only baseline) minimises the heteroscedastic Gaussian NLL: "
    "L_NLL = (1/N) \u03a3_i \u03a3_j [(y_ij \u2212 \u03bc_ij)\u00b2/(2\u03c3_ij\u00b2) "
    "+ (1/2) log \u03c3_ij\u00b2].",
    "纯数据驱动基线代理最小化异方差高斯负对数似然：L_NLL = (1/N)\u03a3_i\u03a3_j "
    "[(y_ij\u2212\u03bc_ij)\u00b2/(2\u03c3_ij\u00b2) + (1/2)log\u03c3_ij\u00b2]。")

add_bilingual(doc,
    "The physics-regularized surrogate (monotone-regularized surrogate) augments L_NLL with a monotonicity "
    "regularization term. For each physically-expected monotone pair (p, q) with expected "
    "sign s_pq \u2208 {+1, \u22121}: L_mono = \u03bb \u03a3_(p,q) max(0, \u03b4 \u2212 "
    "s_pq \u00b7 \u03c1\u0302_pq), where \u03c1\u0302_pq is the batch Spearman rank "
    "correlation. Total loss: L = L_NLL + L_mono. "
    "[Note: monotone pair list, \u03bb and \u03b4 values to be added.]",
    "单调正则化代理在L_NLL基础上附加单调性正则化项：对每个已知单调关系的输入–输出对"
    "(p,q)（期望符号s_pq\u2208{+1,\u22121}），在小批次上计算Spearman秩相关\u03c1\u0302_pq，"
    "施加铰链惩罚：L_mono=\u03bb\u03a3_(p,q)max(0,\u03b4\u2212s_pq\u00b7\u03c1\u0302_pq)，"
    "总损失L=L_NLL+L_mono。【待补充：单调对列表及\u03bb、\u03b4取值】")

add_bilingual(doc,
    "Hyperparameters for both levels were independently optimised using Optuna (40 trials "
    "per level), maximising validation NLL. The split was fixed prior to model selection "
    "(seed 2026) to prevent data leakage.",
    "data-only baseline和monotone-regularized surrogate的超参数均通过Optuna独立优化（每层次40次试验，验证集NLL为目标）。"
    "训练/验证/测试划分在任何模型选择之前固定（种子2026），防止数据泄露。")

# 5.4
add_heading(doc, "5.4  Forward Uncertainty Propagation / 前向不确定性传播", level=2)

add_bilingual(doc,
    "20,000 samples are drawn from the joint prior. For each sample \u03b8_i, the surrogate "
    "returns \u03bc(\u03b8_i). The distribution of predicted means constitutes the epistemic "
    "forward uncertainty propagation. The exceedance probability is: "
    "P_exceed(\u03c4) = (1/N) \u03a3_i \u1d35[\u03bc_stress(\u03b8_i) > \u03c4], "
    "with \u03c4 = 131 MPa as the primary threshold; 110 and 120 MPa in the Appendix.",
    "从联合先验抽取2万个样本，每个样本\u03b8_i通过代理模型得到预测均值\u03bc(\u03b8_i)，"
    "其分布构成认知不确定性前向传播结果。超阈概率：P_exceed(\u03c4)=(1/N)\u03a3_i"
    "\u1d35[\u03bc_stress(\u03b8_i)>\u03c4]，正文主阈值\u03c4=131 MPa，附录另报告110和120 MPa阈值。")

# 5.5
add_heading(doc, "5.5  Global Sensitivity Analysis / 全局敏感性分析", level=2)

add_bilingual(doc,
    "First-order (S\u2081) Sobol indices are estimated using the Jansen estimator with "
    "independent random sampling (N_S = 4096 per matrix). Uncertainty in the indices is "
    "quantified by 50 independent replications with normal-approximation 90% confidence "
    "intervals (\u0305S \u00b1 1.645\u00b7SE). The monotone-regularized surrogate surrogate mean serves as the "
    "function evaluator.",
    "一阶（S\u2081）Sobol指数采用Jansen估计量，基于独立随机采样（每矩阵N_S=4096个样本），"
    "50次独立重复估计，取均值和正态近似置信区间（\u0305S\u00b11.645\u00b7SE，90% CI），"
    "以monotone-regularized surrogate代理均值预测为函数评估器，计算成本仅为直接仿真的极小比例。")

# 5.6
add_heading(doc, "5.6  Observation-Driven Posterior Calibration / 观测驱动后验标定", level=2)

add_bilingual(doc,
    "For 20 benchmark test cases, MCMC with a Gaussian likelihood updates the prior over "
    "4 calibrated parameters (E_intercept, \u03b1_base, \u03b1_slope, \u03bd). The "
    "remaining four parameters (E_slope, T_ref,SS316, k_ref,SS316, k_slope,SS316) are held "
    "at their true test-case input values, reflecting a partial-observation scenario in "
    "which SS316 thermal properties are treated as known from prior material characterisation. "
    "Although k_ref,SS316 is the dominant stress driver (Section 2.3), incorporating it as "
    "a calibration parameter would require independent thermal conductivity measurements; "
    "the present protocol therefore represents a lower bound on achievable posterior "
    "uncertainty. Posterior predictions are validated against the existing HF dataset via "
    "nearest-neighbour lookup (not a true HF rerun). "
    "Each MCMC run generates 1200 posterior samples (acceptance rates 0.40\u20130.53). "
    "P_safe(\u03c4) = 1 \u2212 P_exceed(\u03c4 | data), estimated from posterior predictive samples. "
    "[Note: \u03c3_obs observation noise to be added.]",
    "对20个基准测试案例，以高斯似然函数为条件，MCMC更新4个标定参数（E_intercept, \u03b1_base, "
    "\u03b1_slope, \u03bd）的先验。其余4个参数（E_slope, T_ref,SS316, k_ref,SS316, "
    "k_slope,SS316）固定在对应测试案例的真实输入值，对应于SS316热性质可由材料表征预先知晓的"
    "部分观测场景。尽管k_ref,SS316是应力主控因子，将其纳入标定需要独立热导率测量；当前方案"
    "因此代表后验不确定性的下界。后验验证采用最近邻HF数据集查找方式，而非在后验均值处重新"
    "运行高保真求解器。每次MCMC生成1200个后验样本（接受率0.40\u20130.53）。"
    "P_safe(\u03c4)=1\u2212P_exceed(\u03c4|data)，由后验预测样本估计。"
    "【待补充：\u03c3_obs观测噪声取值】")

hr(doc)

# ── Acknowledgements ──────────────────────────────────────────────────────────
add_heading(doc, "Acknowledgements / 致谢", level=1)

add_para(doc,
    "The authors gratefully acknowledge the support provided by the National Natural Science "
    "Foundation of China [12175138]; Shanghai Rising-Star Program; and LingChuang Research "
    "Project of China National Nuclear Corporation.",
    size=10, space_before=0, space_after=6)

hr(doc)

# ── References placeholder ────────────────────────────────────────────────────
add_heading(doc, "References / 参考文献", level=1)
add_para(doc, "[Reference list to be compiled from sn-article.bbl]",
         size=10, italic=True, color=(120, 120, 120))

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/yinuo/Projects/hpr-claude-project/HPR_surrogate_paper_draft.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
