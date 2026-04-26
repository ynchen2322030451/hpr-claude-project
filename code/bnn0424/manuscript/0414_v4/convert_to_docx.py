#!/usr/bin/env python3
"""Convert draft_paper_0414_v4.txt to a formatted Word document.

Key design decisions:
- Lines within the same paragraph block are joined with a single space
  (no hard line breaks in Word).
- Chinese font (宋体) is set via east-asian font attributes so both
  English (Times New Roman) and Chinese render correctly.
- 【…】 markers are highlighted in red; [TABLE]/[FIG] placeholders in blue.
- Display equations (indented lines with =, ~, Σ, ∫) are rendered as
  Word OMML equation objects (Cambria Math, centered).
- Inline variable names are rendered with italic base + real subscript.
"""

import re
from datetime import datetime
from pathlib import Path
from lxml import etree

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
OMML_NSMAP = {"m": OMML_NS}

HERE = Path(__file__).parent
TXT_PATH = HERE / "draft_paper_0414_v4.txt"
TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
DOCX_PATH = HERE / f"draft_paper_0414_v4_{TIMESTAMP}.docx"
DOCX_LATEST = HERE / "draft_paper_0414_v4.docx"

# ── Font helpers ──────────────────────────────────────────────────────

CN_FONT = "微软雅黑"
EN_FONT = "Times New Roman"


def _set_run_font(run, size=None, bold=None, italic=None, color=None):
    """Set font attributes on a run, including east-asian font for Chinese."""
    rpr = run._element.get_or_add_rPr()
    # English font
    run.font.name = EN_FONT
    # East-asian font for Chinese characters
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def setup_styles(doc):
    """Configure document default style with proper bilingual fonts."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = EN_FONT
    font.size = Pt(11)
    # Set east-asian font in the style XML
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Also set heading styles
    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hfont = hstyle.font
        hfont.name = EN_FONT
        hrpr = hstyle.element.get_or_add_rPr()
        hrFonts = hrpr.find(qn("w:rFonts"))
        if hrFonts is None:
            hrFonts = hstyle.element.makeelement(qn("w:rFonts"), {})
            hrpr.insert(0, hrFonts)
        hrFonts.set(qn("w:eastAsia"), CN_FONT)


# ── OMML equation helpers ────────────────────────────────────────────

def _omml_el(tag, attrib=None, text=None):
    """Create an OMML element with m: namespace."""
    el = etree.SubElement(etree.Element("dummy"), f"{{{OMML_NS}}}{tag}")
    el = etree.Element(f"{{{OMML_NS}}}{tag}", nsmap=OMML_NSMAP)
    if attrib:
        for k, v in attrib.items():
            el.set(f"{{{OMML_NS}}}{k}", v)
    if text is not None:
        el.text = text
    return el


def _omml_run(text, italic=True):
    """Create an OMML <m:r> run with optional italic (math default)."""
    r = etree.Element(f"{{{OMML_NS}}}r", nsmap=OMML_NSMAP)
    if not italic:
        rPr = etree.SubElement(r, f"{{{OMML_NS}}}rPr")
        sty = etree.SubElement(rPr, f"{{{OMML_NS}}}sty")
        sty.set(f"{{{OMML_NS}}}val", "p")
    t = etree.SubElement(r, f"{{{OMML_NS}}}t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r


def _build_omml_paragraph(equation_text):
    """Build an <m:oMathPara> element from plain equation text.

    This creates a display-mode equation using Cambria Math.
    The text is inserted as-is (no LaTeX parsing) but with the
    equation font applied, which gives proper symbol rendering.
    """
    oMathPara = etree.Element(f"{{{OMML_NS}}}oMathPara", nsmap=OMML_NSMAP)
    oMath = etree.SubElement(oMathPara, f"{{{OMML_NS}}}oMath")
    r = etree.SubElement(oMath, f"{{{OMML_NS}}}r")
    rPr = etree.SubElement(r, f"{{{OMML_NS}}}rPr")
    sty = etree.SubElement(rPr, f"{{{OMML_NS}}}sty")
    sty.set(f"{{{OMML_NS}}}val", "pi")
    t = etree.SubElement(r, f"{{{OMML_NS}}}t")
    t.text = equation_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return oMathPara


def _is_display_equation(line):
    """Detect indented display equations (start with 2+ spaces, contain math)."""
    if not line.startswith('  '):
        return False
    stripped = line.strip()
    if not stripped:
        return False
    math_indicators = ('=', '~', '≈', '∫', 'Σ', '∏', '≥', '≤', '∝', '∈')
    return any(c in stripped for c in math_indicators)


def add_display_equation(doc, text):
    """Add a display equation as an OMML equation object (centered, Cambria Math)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    omml = _build_omml_paragraph(text.strip())
    p._element.append(omml)


# ── Paragraph helpers ─────────────────────────────────────────────────

def _join_lines(lines):
    """Join consecutive text lines into a single paragraph string.

    - If two adjacent lines are both Chinese (ending/starting with CJK),
      join without space.
    - Otherwise join with a single space.
    """
    if not lines:
        return ""
    result = lines[0].rstrip()
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            continue
        prev_char = result[-1] if result else ""
        next_char = stripped[0] if stripped else ""
        # If both sides are CJK, no space needed
        if _is_cjk(prev_char) and _is_cjk(next_char):
            result += stripped
        else:
            result += " " + stripped
    return result


def _is_cjk(ch):
    """Check if a character is CJK."""
    if not ch:
        return False
    cp = ord(ch)
    return (
        (0x4E00 <= cp <= 0x9FFF)
        or (0x3400 <= cp <= 0x4DBF)
        or (0xF900 <= cp <= 0xFAFF)
        or (0x2E80 <= cp <= 0x2EFF)
        or (0x3000 <= cp <= 0x303F)
        or (0xFF00 <= cp <= 0xFFEF)
        or (0x2000 <= cp <= 0x206F)  # general punctuation
        or ch in "，。、；：？！""''（）【】—…·"
    )


# Variable names to render as italic base + real subscript in Word.
# Order matters: longer/more-specific patterns first.
# (regex, base, subscript, italic_base)
_VAR_PATTERNS = [
    (r'\bk_?eff\b',       'k', 'eff',       True),
    (r'\bE_intercept\b',  'E', 'intercept', True),
    (r'\bE_slope\b',      'E', 'slope',     True),
    (r'\bk_ref\b',        'k', 'ref',       True),
    (r'\bk_slope\b',      'k', 'slope',     True),
    (r'\bT_ref\b',        'T', 'ref',       True),
    (r'α_base\b',         'α', 'base',      False),
    (r'α_slope\b',        'α', 'slope',     False),
    (r'\balpha_base\b',   'α', 'base',      False),
    (r'\balpha_slope\b',  'α', 'slope',     False),
    (r'σ_prior\b',        'σ', 'prior',     False),
    (r'σ_obs\b',          'σ', 'obs',       False),
    (r'σ_p\b',            'σ', 'p',         False),
    (r'\bN_S\b',          'N', 'S',         True),
    (r'\bS_i\b',          'S', 'i',         True),
    (r'\bV_i\b',          'V', 'i',         True),
    (r'\bS_\{T,i\}\b',   'S', 'T,i',       True),
    (r'S_\{T,i\}',       'S', 'T,i',       True),
    (r'\bPICP₉₀\b',      'PICP', '90',     False),
    (r'\bMPIW₉₀\b',      'MPIW', '90',     False),
    (r'\bR²\b',           'R', '²',         True),
]
_VAR_COMBINED = re.compile(
    '|'.join(f'(?P<v{i}>{pat[0]})' for i, pat in enumerate(_VAR_PATTERNS))
)


def _add_text_with_subscripts(p, text):
    """Add ``text`` to paragraph ``p``, rendering known variable names
    as italic base + real Word subscript (e.g., keff → *k*\u2091ff)."""
    pos = 0
    for m in _VAR_COMBINED.finditer(text):
        start, end = m.span()
        if start > pos:
            run = p.add_run(text[pos:start])
            _set_run_font(run)
        for i, (_, base, sub, italic) in enumerate(_VAR_PATTERNS):
            if m.group(f'v{i}'):
                r_base = p.add_run(base)
                _set_run_font(r_base, italic=italic)
                r_sub = p.add_run(sub)
                _set_run_font(r_sub)
                r_sub.font.subscript = True
                break
        pos = end
    if pos < len(text):
        run = p.add_run(text[pos:])
        _set_run_font(run)


def add_body_paragraph(doc, text):
    """Add a normal paragraph. Highlight 【...】 markers in red bold and
    render known math variables (keff, α_base, E_intercept, …) with
    proper italic base + subscript formatting."""
    p = doc.add_paragraph()
    parts = re.split(r'(【[^】]*】)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('【') and part.endswith('】'):
            run = p.add_run(part)
            _set_run_font(run, color=RGBColor(0xCC, 0x00, 0x00), bold=True)
        else:
            _add_text_with_subscripts(p, part)


def add_placeholder(doc, text):
    """Add [TABLE]/[FIG]/[第X段] placeholder in blue italic."""
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    _set_run_font(run, size=10, italic=True, color=RGBColor(0x00, 0x55, 0xAA))


def add_comment_block(doc, text):
    """Add comment lines (# or %CN:) in gray small italic."""
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    _set_run_font(run, size=9, italic=True, color=RGBColor(0x88, 0x88, 0x88))


def add_table(doc, header_line, rows):
    """Add a markdown-style table."""
    cols = [c.strip() for c in header_line.split('|') if c.strip()]
    n_cols = len(cols)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for i, col in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = col
        for run in cell.paragraphs[0].runs:
            _set_run_font(run, bold=True)
    for r_idx, row_line in enumerate(rows):
        cells_text = [c.strip() for c in row_line.split('|') if c.strip()]
        for c_idx, ct in enumerate(cells_text):
            if c_idx < n_cols:
                table.rows[1 + r_idx].cells[c_idx].text = ct
    doc.add_paragraph()


# ── Main parser ───────────────────────────────────────────────────────

def parse_and_convert(txt_path, docx_path):
    doc = Document()
    setup_styles(doc)

    lines = txt_path.read_text(encoding="utf-8").splitlines()

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "A Posterior-Predictive Bayesian Surrogate for Coupled "
        "Neutronic\u2013Thermal\u2013Structural Analysis in "
        "Heat-Pipe-Cooled Reactors"
    )
    _set_run_font(run, size=14, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "面向热管冷却反应堆耦合中子\u2013热\u2013力分析的"
        "后验预测贝叶斯代理"
    )
    _set_run_font(run2, size=13, bold=True)

    # Version line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        f"Draft v4 (Round-6A prose surgery)  |  Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    _set_run_font(run3, size=9, color=RGBColor(0x99, 0x99, 0x99))
    doc.add_paragraph()  # spacer

    # ── Skip header comment block ──
    i = 0
    while i < len(lines) and (lines[i].startswith('#') or lines[i].strip() == ''):
        i += 1

    # ── State machine ──
    buf = []            # accumulates lines for current paragraph
    table_header = None
    table_rows = []
    in_table = False

    def flush_buf():
        nonlocal buf
        if not buf:
            return
        joined = _join_lines(buf)
        buf = []
        if not joined.strip():
            return
        # Classify
        if joined.startswith('[TABLE') or joined.startswith('[FIG'):
            add_placeholder(doc, joined)
        elif joined.startswith('%CN:') or joined.startswith('#'):
            add_comment_block(doc, joined)
        else:
            add_body_paragraph(doc, joined)

    def flush_table():
        nonlocal table_header, table_rows, in_table
        if table_header and table_rows:
            add_table(doc, table_header, table_rows)
        table_header = None
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ═══ Section heading ═══
        if re.match(r'^={10,}$', stripped):
            flush_buf()
            flush_table()
            if i + 2 < len(lines) and re.match(r'^={10,}$', lines[i + 2].strip()):
                heading = lines[i + 1].strip()
                doc.add_heading(heading, level=1)
                i += 3
                continue
            i += 1
            continue

        # ─── Sub-section heading ───
        if re.match(r'^-{10,}$', stripped):
            flush_buf()
            flush_table()
            if i + 2 < len(lines) and re.match(r'^-{10,}$', lines[i + 2].strip()):
                heading = lines[i + 1].strip()
                doc.add_heading(heading, level=2)
                i += 3
                continue
            i += 1
            continue

        # Markdown table row
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            flush_buf()
            if re.match(r'^\|[-\s|]+\|$', stripped):
                i += 1
                continue
            if not in_table:
                table_header = stripped
                in_table = True
            else:
                table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # %CN: comment lines — collect all consecutive
        if stripped.startswith('%CN:'):
            flush_buf()
            comment_lines = []
            while i < len(lines) and lines[i].strip().startswith('%CN:'):
                comment_lines.append(lines[i].strip())
                i += 1
            add_comment_block(doc, _join_lines(comment_lines))
            continue

        # [TABLE ...] or [FIG ...] placeholders
        if stripped.startswith('[TABLE') or stripped.startswith('[FIG'):
            flush_buf()
            placeholder_parts = [stripped]
            i += 1
            while (
                i < len(lines)
                and lines[i].strip()
                and not lines[i].strip().startswith('[')
                and not re.match(r'^[-=]{10,}$', lines[i].strip())
            ):
                placeholder_parts.append(lines[i].strip())
                i += 1
            add_placeholder(doc, ' '.join(placeholder_parts))
            continue

        # [第X段：...] section hint
        if re.match(r'^\[第.+段[：:]', stripped):
            flush_buf()
            add_placeholder(doc, stripped)
            i += 1
            continue

        # # comment block
        if stripped.startswith('#'):
            flush_buf()
            comment_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('#'):
                comment_lines.append(lines[i].strip())
                i += 1
            add_comment_block(doc, '\n'.join(comment_lines))
            continue

        # Display equation (indented, contains math symbols)
        if _is_display_equation(line):
            flush_buf()
            eq_lines = [line.rstrip()]
            i += 1
            while i < len(lines) and _is_display_equation(lines[i]):
                eq_lines.append(lines[i].rstrip())
                i += 1
            for eq in eq_lines:
                add_display_equation(doc, eq)
            continue

        # Empty line → paragraph break
        if stripped == '':
            flush_buf()
            i += 1
            continue

        # Regular text line → accumulate
        buf.append(line.rstrip())
        i += 1

    flush_buf()
    flush_table()

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    out = parse_and_convert(TXT_PATH, DOCX_PATH)
    print(f"Word document saved to: {out}")
    parse_and_convert(TXT_PATH, DOCX_LATEST)
    print(f"Latest copy saved to:   {DOCX_LATEST}")
