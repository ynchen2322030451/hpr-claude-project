#!/usr/bin/env python3
"""Convert NC-draft txt files to formatted Word documents.

Usage:
    python convert_to_docx.py              # convert all 4 txt files
    python convert_to_docx.py en           # convert EN-only files
    python convert_to_docx.py bilingual    # convert bilingual files

Produces .docx alongside each .txt source.

Features:
- Variable names rendered as italic base + real Word subscript/superscript
  (k_eff, E_intercept, alpha_base, R^2, S_1, S_T, N_S, etc.)
- Scientific notation (10^5, 10^{-2}) rendered with proper superscripts
- CJK support: east-asian font = 微软雅黑, western = Times New Roman
- Section headings: ===...=== (level 1), ---...--- (level 2),
  ###...### (level 1), ### text (level 1)
- Pipe-separated tables -> Word tables with "Light Grid Accent 1" style
  Multi-section tables (sub-headers like "----- Physics-regularized BNN -----")
  are rendered as merged-cell sub-header rows.
- Placeholders [TO GENERATE], [TO ADD], [Source: ...] -> blue italic
- [NOTE: ...] blocks -> gray small italic
- Display equations (2+ space indent with math symbols) -> centered Cambria Math
- Bold markers **text** -> bold
"""

import re
import sys
from pathlib import Path
from lxml import etree

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
OMML_NSMAP = {"m": OMML_NS}

HERE = Path(__file__).parent
_BNN0414 = HERE.parents[1]
FIGURE_DIR = _BNN0414 / "manuscript" / "0414_v4" / "figures"

CN_FONT = "微软雅黑"
EN_FONT = "Times New Roman"

# ── Figure mapping: [Source: ...] path fragments → composed figure PNGs ──
# Maps keywords from [Source: ...] lines to composed figure filenames.
# Order matters: first match wins.
_SOURCE_TO_FIGURE = [
    # SI figures
    ("fig1_workflow",            "fig0_workflow.png"),
    ("sobol_convergence",        "figS1_sobol_convergence.png"),
    ("prior_sensitivity",        "figS2_prior_sensitivity.png"),
    ("noise_sensitivity",        "figS3_noise_sensitivity.png"),
    ("ood_coverage\|ood_epistemic\|ood_calibration", "figS4_ood.png"),
    ("external_baseline_calib\|reliability_mc-dropout\|reliability_deep-ensemble",
                                 "figS5_external_calib.png"),
    ("B1_stress_parity",         "fig2_predictive.png"),
    ("C1_stress_coupling",       "fig3_forward.png"),
    ("E1_prior_posterior",       "fig6_posterior.png"),
    ("E3_posterior_predictive",   "fig6_posterior.png"),
    ("E2_posterior_coverage",     "fig6_posterior.png"),
    ("epi_vs_ale_scatter",       "fig5_physics.png"),
    ("trace_bnn-phy-mono",       "fig6_posterior.png"),
    ("data_efficiency_curve",    "figA3_efficiency.png"),
    ("budget_matched_risk",      "fig7_efficiency.png"),
    ("uncertainty_decomposition", "fig5_physics.png"),
]


def _resolve_figure(source_text: str):
    """Try to find a composed figure PNG matching a [Source: ...] line."""
    for pattern, fname in _SOURCE_TO_FIGURE:
        if re.search(pattern, source_text, re.IGNORECASE):
            fpath = FIGURE_DIR / fname
            if fpath.exists():
                return fpath
    return None

# ── Font helpers ─────────────────────────────────────────────────────


def _set_run_font(run, size=None, bold=None, italic=None, color=None):
    """Set font attributes on a run, including east-asian font for CJK."""
    rpr = run._element.get_or_add_rPr()
    run.font.name = EN_FONT
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def setup_styles(doc):
    """Configure document default style with proper bilingual fonts."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = EN_FONT
    font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hfont = hstyle.font
        hfont.name = EN_FONT
        hrpr = hstyle.element.get_or_add_rPr()
        hrFonts = hrpr.find(qn("w:rFonts"))
        if hrFonts is None:
            hrFonts = hstyle.element.makeelement(qn("w:rFonts"), {})
            hrpr.insert(0, hrFonts)
        hrFonts.set(qn("w:eastAsia"), CN_FONT)


# ── OMML equation helpers ────────────────────────────────────────────


def _build_omml_paragraph(equation_text):
    """Build an <m:oMathPara> element from plain equation text."""
    oMathPara = etree.Element(f"{{{OMML_NS}}}oMathPara", nsmap=OMML_NSMAP)
    oMath = etree.SubElement(oMathPara, f"{{{OMML_NS}}}oMath")
    r = etree.SubElement(oMath, f"{{{OMML_NS}}}r")
    rPr = etree.SubElement(r, f"{{{OMML_NS}}}rPr")
    sty = etree.SubElement(rPr, f"{{{OMML_NS}}}sty")
    sty.set(f"{{{OMML_NS}}}val", "pi")
    t = etree.SubElement(r, f"{{{OMML_NS}}}t")
    t.text = equation_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return oMathPara


def _is_display_equation(line):
    """Detect indented display equations (start with 2+ spaces, contain math)."""
    if not line.startswith('  '):
        return False
    stripped = line.strip()
    if not stripped:
        return False
    math_indicators = ('=', '~', '≈', '∫', 'Σ', '∏', '≥', '≤', '∝', '∈')
    return any(c in stripped for c in math_indicators)


def add_display_equation(doc, text):
    """Add a display equation as an OMML equation object (centered, Cambria Math)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    omml = _build_omml_paragraph(text.strip())
    p._element.append(omml)


# ── CJK detection ───────────────────────────────────────────────────


def _is_cjk(ch):
    """Check if a character is CJK."""
    if not ch:
        return False
    cp = ord(ch)
    return (
        (0x4E00 <= cp <= 0x9FFF)
        or (0x3400 <= cp <= 0x4DBF)
        or (0xF900 <= cp <= 0xFAFF)
        or (0x2E80 <= cp <= 0x2EFF)
        or (0x3000 <= cp <= 0x303F)
        or (0xFF00 <= cp <= 0xFFEF)
        or (0x2000 <= cp <= 0x206F)
        or ch in "，。、；：？！""''（）【】—…·"
    )


def _join_lines(lines):
    """Join consecutive text lines into a single paragraph string.

    - If two adjacent lines are both CJK (ending/starting with CJK),
      join without space.
    - Otherwise join with a single space.
    """
    if not lines:
        return ""
    result = lines[0].rstrip()
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            continue
        prev_char = result[-1] if result else ""
        next_char = stripped[0] if stripped else ""
        if _is_cjk(prev_char) and _is_cjk(next_char):
            result += stripped
        else:
            result += " " + stripped
    return result


# ── Variable subscript / superscript rendering ──────────────────────

# Order matters: longer/more-specific patterns first.
# (regex, base, sub_or_sup, italic_base, is_superscript)
_VAR_PATTERNS = [
    # Subscript variables
    (r'\bk_?eff\b',         'k',    'eff',       True,  False),
    (r'\bE_intercept\b',    'E',    'intercept',  True,  False),
    (r'\bE_slope\b',        'E',    'slope',      True,  False),
    (r'\bk_ref\b',          'k',    'ref',        True,  False),
    (r'\bk_slope\b',        'k',    'slope',      True,  False),
    (r'\bT_ref\b',          'T',    'ref',        True,  False),
    (r'α_base\b',           '\u03b1', 'base',     False, False),
    (r'α_slope\b',          '\u03b1', 'slope',    False, False),
    (r'\balpha_base\b',     '\u03b1', 'base',     False, False),
    (r'\balpha_slope\b',    '\u03b1', 'slope',    False, False),
    (r'σ_prior\b',          '\u03c3', 'prior',    False, False),
    (r'σ_obs\b',            '\u03c3', 'obs',      False, False),
    (r'σ_p\b',              '\u03c3', 'p',        False, False),
    (r'\bN_S\b',            'N',    'S',          True,  False),
    (r'\bS_1\b',            'S',    '1',          True,  False),
    (r'\bS_T\b',            'S',    'T',          True,  False),
    (r'\bS_i\b',            'S',    'i',          True,  False),
    (r'\bV_i\b',            'V',    'i',          True,  False),
    (r'S_\{T,i\}',         'S',    'T,i',        True,  False),
    (r'\bMW_e\b',           'MW',   'e',          False, False),
    (r'\bPICP₉₀\b',        'PICP', '90',         False, False),
    (r'\bMPIW₉₀\b',        'MPIW', '90',         False, False),
    # Superscript: R^2
    (r'\bR\^2\b',           'R',    '2',          True,  True),
    (r'\bR²\b',             'R',    '2',          True,  True),
]
_VAR_COMBINED = re.compile(
    '|'.join(f'(?P<v{i}>{pat[0]})' for i, pat in enumerate(_VAR_PATTERNS))
)

# Scientific notation: 10^5, 10^{-2}, 10^{-5}, etc.
# Also handles x 10^... and similar
_SCI_NOTATION = re.compile(
    r'10\^'
    r'(?:\{([^}]+)\}|(\d+))'
)

# Bold markers: **text**
_BOLD_PATTERN = re.compile(r'\*\*([^*]+)\*\*')


def _add_text_with_formatting(p, text):
    """Add text to paragraph with variable subscripts/superscripts,
    scientific notation superscripts, and bold markers."""
    # Build a combined pattern for all special formatting
    # We process in order: find all matches, sort by position, render segments

    matches = []

    # Variable subscripts/superscripts
    for m in _VAR_COMBINED.finditer(text):
        for i, (_, base, sub, italic, is_sup) in enumerate(_VAR_PATTERNS):
            if m.group(f'v{i}'):
                matches.append((m.start(), m.end(), 'var', i))
                break

    # Scientific notation
    for m in _SCI_NOTATION.finditer(text):
        matches.append((m.start(), m.end(), 'sci', m))

    # Bold markers
    for m in _BOLD_PATTERN.finditer(text):
        matches.append((m.start(), m.end(), 'bold', m))

    # Sort by start position; if overlapping, keep the earlier/longer one
    matches.sort(key=lambda x: (x[0], -x[1]))

    # Remove overlapping matches (keep first)
    filtered = []
    last_end = 0
    for start, end, kind, data in matches:
        if start >= last_end:
            filtered.append((start, end, kind, data))
            last_end = end
    matches = filtered

    pos = 0
    for start, end, kind, data in matches:
        # Plain text before this match
        if start > pos:
            run = p.add_run(text[pos:start])
            _set_run_font(run)

        if kind == 'var':
            i = data
            _, base, sub, italic, is_sup = _VAR_PATTERNS[i]
            r_base = p.add_run(base)
            _set_run_font(r_base, italic=italic)
            r_sub = p.add_run(sub)
            _set_run_font(r_sub)
            if is_sup:
                r_sub.font.superscript = True
            else:
                r_sub.font.subscript = True

        elif kind == 'sci':
            m = data
            # Render "10" as normal text, exponent as superscript
            r_base = p.add_run("10")
            _set_run_font(r_base)
            exponent = m.group(1) if m.group(1) else m.group(2)
            r_sup = p.add_run(exponent)
            _set_run_font(r_sup)
            r_sup.font.superscript = True

        elif kind == 'bold':
            m = data
            r_bold = p.add_run(m.group(1))
            _set_run_font(r_bold, bold=True)

        pos = end

    # Trailing text
    if pos < len(text):
        run = p.add_run(text[pos:])
        _set_run_font(run)


# ── Paragraph helpers ────────────────────────────────────────────────


def add_body_paragraph(doc, text):
    """Add a normal paragraph with rich formatting:
    - [TO ...] / [Source: ...] placeholders in blue italic
    - [NOTE: ...] in gray italic
    - Variable subscripts/superscripts
    - Scientific notation
    - Bold markers
    """
    p = doc.add_paragraph()
    # Split on placeholder patterns and [NOTE: ...] blocks
    parts = re.split(
        r'(\[(?:TO [A-Z][^\]]*|Source:[^\]]*|NOTE:[^\]]*)\])',
        text
    )
    for part in parts:
        if not part:
            continue
        if re.match(r'\[TO [A-Z]', part) or part.startswith('[Source:'):
            run = p.add_run(part)
            _set_run_font(run, italic=True, color=RGBColor(0x00, 0x55, 0xAA))
        elif part.startswith('[NOTE:'):
            run = p.add_run(part)
            _set_run_font(run, size=9, italic=True, color=RGBColor(0x88, 0x88, 0x88))
        else:
            _add_text_with_formatting(p, part)


def add_placeholder(doc, text, embed_figures=True):
    """Add [SOURCE]/[TO ...] placeholder in blue italic.

    If embed_figures is True and the placeholder is a [Source: ...] line
    with a matching composed figure, embed the figure image inline.
    """
    stripped = text.strip()
    if embed_figures and stripped.startswith("[Source:"):
        fig_path = _resolve_figure(stripped)
        if fig_path:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(fig_path), width=Inches(6.0))
            # Also add a small source note below
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(stripped)
            _set_run_font(run2, size=8, italic=True,
                          color=RGBColor(0xAA, 0xAA, 0xAA))
            return
    p = doc.add_paragraph()
    run = p.add_run(stripped)
    _set_run_font(run, size=10, italic=True, color=RGBColor(0x00, 0x55, 0xAA))


def add_note_block(doc, text):
    """Add [NOTE: ...] block in gray small italic."""
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    _set_run_font(run, size=9, italic=True, color=RGBColor(0x88, 0x88, 0x88))


def add_table_block(doc, header_line, rows):
    """Add a markdown-style table to the document.

    Handles multi-section tables: rows matching
    "----- Sub-header text -----" are rendered as merged sub-header rows.
    """
    cols = [c.strip() for c in header_line.split('|') if c.strip()]
    n_cols = len(cols)
    if n_cols == 0:
        return

    # Pre-scan rows to count actual data rows vs sub-header rows
    row_specs = []  # (type, content) where type is 'data' or 'subheader'
    for row_line in rows:
        stripped = row_line.strip()
        # Detect sub-header rows like "----- Physics-regularized BNN -----"
        m = re.match(r'^-{3,}\s+(.+?)\s+-{3,}$', stripped)
        if m:
            row_specs.append(('subheader', m.group(1)))
        else:
            row_specs.append(('data', row_line))

    total_rows = 1 + len(row_specs)  # header + all rows
    table = doc.add_table(rows=total_rows, cols=n_cols)
    try:
        table.style = 'Light Grid Accent 1'
    except KeyError:
        pass  # fall back to default if style not available

    # Fill header row
    for i, col in enumerate(cols):
        cell = table.rows[0].cells[i]
        # Clear default paragraph and use our formatting
        cell.paragraphs[0].clear()
        _add_text_with_formatting(cell.paragraphs[0], col)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Fill data/sub-header rows
    for r_idx, (rtype, content) in enumerate(row_specs):
        if rtype == 'subheader':
            # Merge all cells and write sub-header text
            row = table.rows[1 + r_idx]
            # Merge cells
            if n_cols > 1:
                for ci in range(1, n_cols):
                    row.cells[0].merge(row.cells[ci])
            cell = row.cells[0]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(content)
            _set_run_font(run, bold=True, italic=True)
        else:
            cells_text = [c.strip() for c in content.split('|') if c.strip()]
            for c_idx, ct in enumerate(cells_text):
                if c_idx < n_cols:
                    cell = table.rows[1 + r_idx].cells[c_idx]
                    cell.paragraphs[0].clear()
                    _add_text_with_formatting(cell.paragraphs[0], ct)

    doc.add_paragraph()  # spacer after table


# ── Main parser ──────────────────────────────────────────────────────


def parse_and_convert(txt_path, docx_path, *, append_figures=False):
    doc = Document()
    setup_styles(doc)

    lines = txt_path.read_text(encoding="utf-8").splitlines()

    i = 0
    buf = []
    in_table = False
    table_header = None
    table_rows = []

    def flush_buf():
        nonlocal buf
        if not buf:
            return
        joined = _join_lines(buf)
        buf = []
        if not joined.strip():
            return
        stripped = joined.strip()
        if stripped.startswith('[Source:') or stripped.startswith('[TO '):
            add_placeholder(doc, joined)
        elif stripped.startswith('[NOTE:'):
            add_note_block(doc, joined)
        else:
            add_body_paragraph(doc, joined)

    def flush_table():
        nonlocal table_header, table_rows, in_table
        if table_header and table_rows:
            add_table_block(doc, table_header, table_rows)
        table_header = None
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ═══ Major heading (===...===)
        # May contain multiple lines between delimiters (e.g. bilingual titles)
        if re.match(r'^={10,}$', stripped):
            flush_buf()
            flush_table()
            j = i + 1
            heading_lines = []
            while j < len(lines) and not re.match(r'^={10,}$', lines[j].strip()):
                if lines[j].strip():
                    heading_lines.append(lines[j].strip())
                j += 1
            if heading_lines and j < len(lines) and re.match(r'^={10,}$', lines[j].strip()):
                heading = _join_lines(heading_lines)
                doc.add_heading(heading, level=1)
                i = j + 1
                continue
            i += 1
            continue

        # When inside a table, a line of dashes is a table separator -- skip
        if in_table and re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # ─── Sub heading (---...---)
        # Two formats:
        #   (a) Sandwich:  ---...--- / text (1-3 lines, no blank) / ---...---
        #       Used for supplementary figure/table captions
        #   (b) Underline: text / ---...---
        #       Used for manuscript sub-sections
        if re.match(r'^-{10,}$', stripped):
            # (a) Sandwich: look ahead for closing ---, but only within
            # a few lines (max 5) and no blank lines allowed
            j = i + 1
            heading_lines = []
            is_sandwich = False
            while j < len(lines) and (j - i) <= 5:
                lj = lines[j].strip()
                if re.match(r'^-{10,}$', lj):
                    if heading_lines:
                        is_sandwich = True
                    break
                if not lj:
                    break  # blank line breaks the sandwich
                heading_lines.append(lj)
                j += 1
            if is_sandwich:
                flush_buf()
                heading = _join_lines(heading_lines)
                doc.add_heading(heading, level=2)
                i = j + 1
                continue
            # (b) Underline: previous buffered text is the heading,
            # BUT only if it does not look like a table header row
            if buf:
                joined_buf = _join_lines(buf)
                buf_cells = [c.strip() for c in joined_buf.split('|') if c.strip()]
                if len(buf_cells) >= 3:
                    # Looks like a table header row, not a heading.
                    # Start a table with this as header.
                    table_header = joined_buf
                    in_table = True
                    buf.clear()
                    i += 1
                    continue
                heading_text = joined_buf
                buf.clear()
                doc.add_heading(heading_text, level=2)
                i += 1
                continue
            # Otherwise just a separator line -- skip
            flush_buf()
            i += 1
            continue

        # ####...#### section divider (used in supplementary info)
        if re.match(r'^#{10,}$', stripped):
            flush_buf()
            flush_table()
            j = i + 1
            heading_lines = []
            while j < len(lines) and not re.match(r'^#{10,}$', lines[j].strip()):
                if lines[j].strip():
                    heading_lines.append(lines[j].strip())
                j += 1
            if heading_lines and j < len(lines) and re.match(r'^#{10,}$', lines[j].strip()):
                heading = _join_lines(heading_lines)
                doc.add_heading(heading, level=1)
                i = j + 1
                continue
            i += 1
            continue

        # ### heading (Markdown-style)
        m_heading = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if m_heading:
            flush_buf()
            flush_table()
            level = len(m_heading.group(1))
            doc.add_heading(m_heading.group(2), level=level)
            i += 1
            continue

        # Subsection headings like "A.1. Title" or "B.2. Title" followed
        # by a line of dashes  ---...---
        # These are level-3 headings in the supplementary notes.
        # May span multiple lines (e.g. bilingual D.1. ... / Chinese)
        if re.match(r'^[A-Z]\.\d+\.?\s+', stripped):
            flush_buf()
            flush_table()
            # Collect heading lines until we hit a ---...--- line
            heading_parts = [stripped]
            j = i + 1
            while (j < len(lines)
                   and lines[j].strip()
                   and not re.match(r'^-{3,}$', lines[j].strip())):
                heading_parts.append(lines[j].strip())
                j += 1
            if j < len(lines) and re.match(r'^-{3,}$', lines[j].strip()):
                heading = _join_lines(heading_parts)
                doc.add_heading(heading, level=3)
                i = j + 1  # skip the dashes line too
                continue
            else:
                # No dashes found -- not a heading, treat as regular text
                buf.extend(heading_parts)
                i = j
                continue

        # Table sub-header lines inside a table context:
        # e.g. " ----- Physics-regularized BNN -----" (no | chars)
        if in_table and re.match(r'^-{3,}\s+(.+?)\s+-{3,}$', stripped):
            table_rows.append(stripped)
            i += 1
            continue

        # Table row detection: contains | and not a placeholder/note
        # Require at least 3 non-empty cells to distinguish from inline math
        # notation like "p(theta | y_obs)" which has | but is not a table.
        _is_table_row = False
        if ('|' in stripped
                and not stripped.startswith('[')
                and not stripped.startswith('#')):
            # Skip pure separator lines like "-------" or "------|------"
            if re.match(r'^[-\s|]+$', stripped):
                i += 1
                continue
            # Count non-empty cells
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            # Starting a new table requires >= 3 cells (to avoid false
            # positives from inline math like "p(theta | y_obs)").
            # Continuing an existing table allows >= 1 cell (for sparse
            # continuation rows like "intercept | | | |").
            if len(cells) >= 3 or (in_table and len(cells) >= 1):
                _is_table_row = True
        if _is_table_row:
            flush_buf()
            if not in_table:
                table_header = stripped
                in_table = True
            else:
                table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # [Source: ...] placeholder (possibly multi-line)
        if stripped.startswith('[Source:'):
            flush_buf()
            placeholder_parts = [stripped]
            i += 1
            # Collect continuation lines (indented sources)
            while (i < len(lines)
                   and lines[i].strip()
                   and (lines[i].strip().startswith('figures/')
                        or lines[i].strip().startswith('results/')
                        or lines[i].strip().startswith('posterior/')
                        or lines[i].strip().endswith('.png')
                        or lines[i].strip().endswith('.csv')
                        or lines[i].strip().endswith('.pdf')
                        or lines[i].strip().endswith('.jpg'))):
                placeholder_parts.append(lines[i].strip())
                i += 1
            add_placeholder(doc, ' '.join(placeholder_parts))
            continue

        # [TO ...] placeholder (possibly multi-line)
        if stripped.startswith('[TO '):
            flush_buf()
            placeholder_parts = [stripped]
            i += 1
            while (i < len(lines)
                   and lines[i].strip()
                   and not lines[i].strip().startswith('[')
                   and not re.match(r'^[-=]{10,}$', lines[i].strip())
                   and not re.match(r'^#{10,}$', lines[i].strip())):
                placeholder_parts.append(lines[i].strip())
                i += 1
            add_placeholder(doc, ' '.join(placeholder_parts))
            continue

        # [NOTE: ...] block (possibly multi-line)
        if stripped.startswith('[NOTE:'):
            flush_buf()
            note_parts = [stripped]
            i += 1
            while (i < len(lines)
                   and lines[i].strip()
                   and not lines[i].strip().startswith('[')
                   and not re.match(r'^[-=]{10,}$', lines[i].strip())
                   and not re.match(r'^#{10,}$', lines[i].strip())):
                note_parts.append(lines[i].strip())
                i += 1
            add_note_block(doc, ' '.join(note_parts))
            continue

        # Display equation (indented, contains math symbols)
        if _is_display_equation(line):
            flush_buf()
            eq_lines = [line.rstrip()]
            i += 1
            while i < len(lines) and _is_display_equation(lines[i]):
                eq_lines.append(lines[i].rstrip())
                i += 1
            for eq in eq_lines:
                add_display_equation(doc, eq)
            continue

        # Empty line -> paragraph break
        if stripped == '':
            flush_buf()
            i += 1
            continue

        # Regular text line -> accumulate
        buf.append(line.rstrip())
        i += 1

    flush_buf()
    flush_table()

    if append_figures:
        add_figure_legends(doc)

    doc.save(str(docx_path))
    return docx_path


# ── Figure legends (main manuscript) ─────────────────────────────────

_MAIN_FIGURES = [
    ("fig0_workflow.png",   "Figure 1",
     "Overview of the probabilistic surrogate modelling workflow."),
    ("fig1_accuracy.png",   "Figure 2",
     "Prediction accuracy and calibration quality."),
    ("fig2_predictive.png", "Figure 3",
     "Predictive distributions for three representative outputs."),
    ("fig3_forward.png",    "Figure 4",
     "Forward uncertainty quantification and risk assessment."),
    ("fig4_sobol.png",      "Figure 5",
     "Sobol sensitivity analysis for coupled stress and criticality."),
    ("fig5_physics.png",    "Figure 6",
     "Physics consistency: monotonicity and uncertainty decomposition."),
    ("fig6_posterior.png",  "Figure 7",
     "Posterior calibration: MCMC diagnostics and parameter contraction."),
    ("fig7_efficiency.png", "Figure 8",
     "Computational efficiency: speedup and data efficiency."),
]


def add_figure_legends(doc):
    """Append a Figure Legends section with embedded images."""
    doc.add_page_break()
    doc.add_heading("Figure Legends", level=1)

    for fname, label, caption in _MAIN_FIGURES:
        fpath = FIGURE_DIR / fname
        if not fpath.exists():
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: [figure file not found: {fname}]")
            _set_run_font(run, italic=True, color=RGBColor(0xCC, 0x00, 0x00))
            continue

        # Figure image
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(str(fpath), width=Inches(6.0))

        # Caption
        p_cap = doc.add_paragraph()
        run_label = p_cap.add_run(f"{label}. ")
        _set_run_font(run_label, bold=True)
        run_cap = p_cap.add_run(caption)
        _set_run_font(run_cap)

        # Spacer
        doc.add_paragraph()


# ── Entry point ──────────────────────────────────────────────────────

# (txt_path, docx_path, append_figures)
FILES = {
    "en": [
        (HERE / "en" / "manuscript_en.txt",
         HERE / "en" / "manuscript_en.docx", True),
        (HERE / "en" / "supplementary_information_en.txt",
         HERE / "en" / "supplementary_information_en.docx", False),
    ],
    "bilingual": [
        (HERE / "bilingual" / "manuscript_bilingual.txt",
         HERE / "bilingual" / "manuscript_bilingual.docx", True),
        (HERE / "bilingual" / "supplementary_information_bilingual.txt",
         HERE / "bilingual" / "supplementary_information_bilingual.docx", False),
    ],
}

if __name__ == "__main__":
    targets = sys.argv[1:] if len(sys.argv) > 1 else ["en", "bilingual"]
    for group in targets:
        if group not in FILES:
            print(f"Unknown group: {group}. Use 'en' or 'bilingual'.")
            continue
        for txt, docx, figs in FILES[group]:
            if not txt.exists():
                print(f"  SKIP (not found): {txt}")
                continue
            out = parse_and_convert(txt, docx, append_figures=figs)
            print(f"  OK: {out}")
    print("Done.")
