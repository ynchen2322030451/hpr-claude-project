#!/usr/bin/env python3
"""Convert draft_paper_0411_v3.txt to a formatted Word document.

Key design decisions:
- Lines within the same paragraph block are joined with a single space
  (no hard line breaks in Word).
- Chinese font (宋体) is set via east-asian font attributes so both
  English (Times New Roman) and Chinese render correctly.
- 【…】 markers are highlighted in red; [TABLE]/[FIG] placeholders in blue.
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

HERE = Path(__file__).parent
TXT_PATH = HERE / "draft_paper_0411_v3.txt"
TIMESTAMP = datetime.now().strftime("%Y%m%d_%H%M%S")
DOCX_PATH = HERE / f"draft_paper_0411_v3_{TIMESTAMP}.docx"
DOCX_LATEST = HERE / "draft_paper_0411_v3.docx"

# ── Font helpers ──────────────────────────────────────────────────────

CN_FONT = "微软雅黑"
EN_FONT = "Times New Roman"


def _set_run_font(run, size=None, bold=None, italic=None, color=None):
    """Set font attributes on a run, including east-asian font for Chinese."""
    rpr = run._element.get_or_add_rPr()
    # English font
    run.font.name = EN_FONT
    # East-asian font for Chinese characters
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def setup_styles(doc):
    """Configure document default style with proper bilingual fonts."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = EN_FONT
    font.size = Pt(11)
    # Set east-asian font in the style XML
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = style.element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5

    # Also set heading styles
    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hfont = hstyle.font
        hfont.name = EN_FONT
        hrpr = hstyle.element.get_or_add_rPr()
        hrFonts = hrpr.find(qn("w:rFonts"))
        if hrFonts is None:
            hrFonts = hstyle.element.makeelement(qn("w:rFonts"), {})
            hrpr.insert(0, hrFonts)
        hrFonts.set(qn("w:eastAsia"), CN_FONT)


# ── Paragraph helpers ─────────────────────────────────────────────────

def _join_lines(lines):
    """Join consecutive text lines into a single paragraph string.

    - If two adjacent lines are both Chinese (ending/starting with CJK),
      join without space.
    - Otherwise join with a single space.
    """
    if not lines:
        return ""
    result = lines[0].rstrip()
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            continue
        prev_char = result[-1] if result else ""
        next_char = stripped[0] if stripped else ""
        # If both sides are CJK, no space needed
        if _is_cjk(prev_char) and _is_cjk(next_char):
            result += stripped
        else:
            result += " " + stripped
    return result


def _is_cjk(ch):
    """Check if a character is CJK."""
    if not ch:
        return False
    cp = ord(ch)
    return (
        (0x4E00 <= cp <= 0x9FFF)
        or (0x3400 <= cp <= 0x4DBF)
        or (0xF900 <= cp <= 0xFAFF)
        or (0x2E80 <= cp <= 0x2EFF)
        or (0x3000 <= cp <= 0x303F)
        or (0xFF00 <= cp <= 0xFFEF)
        or (0x2000 <= cp <= 0x206F)  # general punctuation
        or ch in "，。、；：？！""''（）【】—…·"
    )


def add_body_paragraph(doc, text):
    """Add a normal paragraph. Highlight 【...】 markers in red bold."""
    p = doc.add_paragraph()
    parts = re.split(r'(【[^】]*】)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('【') and part.endswith('】'):
            run = p.add_run(part)
            _set_run_font(run, color=RGBColor(0xCC, 0x00, 0x00), bold=True)
        else:
            run = p.add_run(part)
            _set_run_font(run)


def add_placeholder(doc, text):
    """Add [TABLE]/[FIG]/[第X段] placeholder in blue italic."""
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    _set_run_font(run, size=10, italic=True, color=RGBColor(0x00, 0x55, 0xAA))


def add_comment_block(doc, text):
    """Add comment lines (# or %CN:) in gray small italic."""
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    _set_run_font(run, size=9, italic=True, color=RGBColor(0x88, 0x88, 0x88))


def add_table(doc, header_line, rows):
    """Add a markdown-style table."""
    cols = [c.strip() for c in header_line.split('|') if c.strip()]
    n_cols = len(cols)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for i, col in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = col
        for run in cell.paragraphs[0].runs:
            _set_run_font(run, bold=True)
    for r_idx, row_line in enumerate(rows):
        cells_text = [c.strip() for c in row_line.split('|') if c.strip()]
        for c_idx, ct in enumerate(cells_text):
            if c_idx < n_cols:
                table.rows[1 + r_idx].cells[c_idx].text = ct
    doc.add_paragraph()


# ── Main parser ───────────────────────────────────────────────────────

def parse_and_convert(txt_path, docx_path):
    doc = Document()
    setup_styles(doc)

    lines = txt_path.read_text(encoding="utf-8").splitlines()

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Probabilistic Neural Surrogates for Uncertainty-to-Risk Analysis "
        "in Coupled Multi-Physics Systems: Application to a Heat-Pipe-Cooled Reactor"
    )
    _set_run_font(run, size=14, bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "概率神经代理模型在耦合多物理场系统不确定性\u2013风险分析中的应用"
        "\u2014\u2014以热管冷却反应堆为例"
    )
    _set_run_font(run2, size=13, bold=True)

    # Version line
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(
        f"Draft v3  |  Generated {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    _set_run_font(run3, size=9, color=RGBColor(0x99, 0x99, 0x99))
    doc.add_paragraph()  # spacer

    # ── Skip header comment block ──
    i = 0
    while i < len(lines) and (lines[i].startswith('#') or lines[i].strip() == ''):
        i += 1

    # ── State machine ──
    buf = []            # accumulates lines for current paragraph
    table_header = None
    table_rows = []
    in_table = False

    def flush_buf():
        nonlocal buf
        if not buf:
            return
        joined = _join_lines(buf)
        buf = []
        if not joined.strip():
            return
        # Classify
        if joined.startswith('[TABLE') or joined.startswith('[FIG'):
            add_placeholder(doc, joined)
        elif joined.startswith('%CN:') or joined.startswith('#'):
            add_comment_block(doc, joined)
        else:
            add_body_paragraph(doc, joined)

    def flush_table():
        nonlocal table_header, table_rows, in_table
        if table_header and table_rows:
            add_table(doc, table_header, table_rows)
        table_header = None
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ═══ Section heading ═══
        if re.match(r'^={10,}$', stripped):
            flush_buf()
            flush_table()
            if i + 2 < len(lines) and re.match(r'^={10,}$', lines[i + 2].strip()):
                heading = lines[i + 1].strip()
                doc.add_heading(heading, level=1)
                i += 3
                continue
            i += 1
            continue

        # ─── Sub-section heading ───
        if re.match(r'^-{10,}$', stripped):
            flush_buf()
            flush_table()
            if i + 2 < len(lines) and re.match(r'^-{10,}$', lines[i + 2].strip()):
                heading = lines[i + 1].strip()
                doc.add_heading(heading, level=2)
                i += 3
                continue
            i += 1
            continue

        # Markdown table row
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            flush_buf()
            if re.match(r'^\|[-\s|]+\|$', stripped):
                i += 1
                continue
            if not in_table:
                table_header = stripped
                in_table = True
            else:
                table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # %CN: comment lines — collect all consecutive
        if stripped.startswith('%CN:'):
            flush_buf()
            comment_lines = []
            while i < len(lines) and lines[i].strip().startswith('%CN:'):
                comment_lines.append(lines[i].strip())
                i += 1
            add_comment_block(doc, _join_lines(comment_lines))
            continue

        # [TABLE ...] or [FIG ...] placeholders
        if stripped.startswith('[TABLE') or stripped.startswith('[FIG'):
            flush_buf()
            placeholder_parts = [stripped]
            i += 1
            while (
                i < len(lines)
                and lines[i].strip()
                and not lines[i].strip().startswith('[')
                and not re.match(r'^[-=]{10,}$', lines[i].strip())
            ):
                placeholder_parts.append(lines[i].strip())
                i += 1
            add_placeholder(doc, ' '.join(placeholder_parts))
            continue

        # [第X段：...] section hint
        if re.match(r'^\[第.+段[：:]', stripped):
            flush_buf()
            add_placeholder(doc, stripped)
            i += 1
            continue

        # # comment block
        if stripped.startswith('#'):
            flush_buf()
            comment_lines = [stripped]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('#'):
                comment_lines.append(lines[i].strip())
                i += 1
            add_comment_block(doc, '\n'.join(comment_lines))
            continue

        # Empty line → paragraph break
        if stripped == '':
            flush_buf()
            i += 1
            continue

        # Regular text line → accumulate
        buf.append(line.rstrip())
        i += 1

    flush_buf()
    flush_table()

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    out = parse_and_convert(TXT_PATH, DOCX_PATH)
    print(f"Word document saved to: {out}")
    parse_and_convert(TXT_PATH, DOCX_LATEST)
    print(f"Latest copy saved to:   {DOCX_LATEST}")
